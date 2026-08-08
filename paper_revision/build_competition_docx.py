from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REVISION = ROOT / "paper_revision"
ARTIFACTS = REVISION / "artifacts"
FIGURES = REVISION / "assets" / "figures"
SCREENSHOTS = REVISION / "assets" / "screenshots"

DESIGN = {
    "east_asia_font": "宋体",
    "latin_font": "Times New Roman",
    "primary": "355C6B",
    "secondary": "7B9AA4",
    "accent": "B58A56",
    "header_fill": "E6EEF1",
    "soft_fill": "F4F7F8",
    "warm_fill": "F4EFE8",
    "text": "222222",
    "muted": "666666",
}

_UNSUPPORTED_ASSERTIONS = [
    re.compile(r"AUC\s*[=:：]\s*0\.91", re.I),
    re.compile(r"Kappa\s*[=:：]\s*0\.86", re.I),
    re.compile(r"Pearson\s*r\s*[=:：]\s*0\.72", re.I),
    re.compile(r"纠正约?6亿元"),
    re.compile(r"6,?3\d{2}条/秒"),
]


def assert_no_unsupported_assertions(text: str) -> None:
    """Reject legacy numeric claims that have no reproducible evidence in the workspace."""
    if "未设置" in text or "不报告" in text or "删除" in text:
        return
    for pattern in _UNSUPPORTED_ASSERTIONS:
        if pattern.search(text):
            raise ValueError(f"发现未获证据支持的断言：{pattern.pattern}")


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = DESIGN["latin_font"]
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), DESIGN["latin_font"])
    rfonts.set(qn("w:hAnsi"), DESIGN["latin_font"])
    rfonts.set(qn("w:eastAsia"), DESIGN["east_asia_font"])
    rfonts.set(qn("w:cs"), DESIGN["latin_font"])
    return run


def shade_cell(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, top="12", header_bottom="8", bottom="12"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is not None:
        tblpr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge, val, size in (
        ("top", "single", top),
        ("bottom", "single", bottom),
        ("insideH", "nil", "0"),
    ):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), DESIGN["primary"])
        borders.append(node)
    tblpr.append(borders)

    header = table.rows[0]
    for cell in header.cells:
        tcpr = cell._tc.get_or_add_tcPr()
        tcborders = tcpr.find(qn("w:tcBorders"))
        if tcborders is None:
            tcborders = OxmlElement("w:tcBorders")
            tcpr.append(tcborders)
        bottom_node = OxmlElement("w:bottom")
        bottom_node.set(qn("w:val"), "single")
        bottom_node.set(qn("w:sz"), header_bottom)
        bottom_node.set(qn("w:color"), DESIGN["secondary"])
        tcborders.append(bottom_node)


def add_three_line_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths_cm: Sequence[float] | None = None,
    font_size: float = 9.5,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        shade_cell(cell, DESIGN["header_fill"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths_cm:
            cell.width = Cm(widths_cm[index])
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                set_run_font(run, font_size, True, DESIGN["primary"])
        set_cell_margins(cell)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths_cm:
                cell.width = Cm(widths_cm[index])
            if len(table.rows) % 2 == 1:
                shade_cell(cell, "F8FAFA")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, font_size, False, DESIGN["text"])
            set_cell_margins(cell)
    set_table_borders(table)
    return table


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    return run


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— ")
    set_run_font(run, 9, color=DESIGN["muted"])
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" —")
    set_run_font(run, 9, color=DESIGN["muted"])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(28)
    section.right_margin = Mm(24)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal = doc.styles["Normal"]
    normal.font.name = DESIGN["latin_font"]
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DESIGN["east_asia_font"])
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = {
        "Title": (24, WD_ALIGN_PARAGRAPH.CENTER, DESIGN["primary"]),
        "Subtitle": (14, WD_ALIGN_PARAGRAPH.CENTER, DESIGN["muted"]),
        "Heading 1": (16, WD_ALIGN_PARAGRAPH.CENTER, DESIGN["primary"]),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT, DESIGN["primary"]),
        "Heading 3": (12, WD_ALIGN_PARAGRAPH.LEFT, DESIGN["secondary"]),
    }
    for name, (size, align, color) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = DESIGN["latin_font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DESIGN["east_asia_font"])
        style.paragraph_format.alignment = align
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(14 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(8)
    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    header = section.header.paragraphs[0]
    header.text = "体融识界·SportFusion｜体育业务边界识别与相对规模测算"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, 9, False, DESIGN["muted"])
    add_page_number(section.footer.paragraphs[0])

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, no_indent=False, align=None):
    assert_no_unsupported_assertions(text)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(0 if no_indent else 24)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 12, True, DESIGN["primary"])
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 12, False, DESIGN["text"])
    else:
        r = p.add_run(text)
        set_run_font(r, 12, False, DESIGN["text"])
    return p


def add_bullet(doc: Document, label: str, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(label), 12, True, DESIGN["primary"])
    set_run_font(p.add_run(text), 12, False, DESIGN["text"])
    return p


def add_formula(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run(text), 12, True, DESIGN["primary"])
    return p


def add_caption(doc: Document, text: str, source: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = bool(source)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    set_run_font(p.add_run(text), 10.5, True, DESIGN["primary"])
    if source:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.first_line_indent = Pt(0)
        s.paragraph_format.space_after = Pt(5)
        set_run_font(s.add_run(f"资料来源：{source}"), 9, False, DESIGN["muted"])


def add_figure(doc: Document, path: Path, caption: str, source: str, width_cm: float = 15.2):
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption, source)


def add_table_caption(doc: Document, text: str, source: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(text), 10.5, True, DESIGN["primary"])
    if source:
        p.add_run()


def add_callout(doc: Document, title: str, text: str, fill: str | None = None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill or DESIGN["soft_fill"])
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "start", "end"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6" if edge == "start" else "0")
        node.set(qn("w:color"), DESIGN["accent"] if edge == "start" else "FFFFFF")
        borders.append(node)
    tcpr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.35
    set_run_font(p.add_run(title), 11, True, DESIGN["primary"])
    set_run_font(p.add_run(text), 11, False, DESIGN["text"])
    set_cell_margins(cell, top=140, start=180, bottom=140, end=140)
    return table


def cover_page(doc: Document, audit: dict):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("体融识界·SportFusion"), 27, True, DESIGN["primary"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("基于NLP文本识别与多维度加权的"), 18, True, DESIGN["text"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("多元经营企业体育业务边界识别与产业相对规模测算"), 18, True, DESIGN["text"])

    accent = doc.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent.cell(0, 0).text = "项目报告｜数据复核修订版"
    shade_cell(accent.cell(0, 0), DESIGN["primary"])
    for run in accent.cell(0, 0).paragraphs[0].runs:
        set_run_font(run, 12, True, "FFFFFF")
    accent.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_margins(accent.cell(0, 0), top=110, start=240, bottom=110, end=240)

    for _ in range(5):
        doc.add_paragraph()
    facts = [
        ["样本记录", f"{audit['snapshot']['total_enterprises']:,} 家"],
        ["正式批次", audit["snapshot"]["formal_batch"]],
        ["核心口径", "相对产出指数（非货币规模）"],
        ["数据原则", "可复算、可追溯、无证据不入结论"],
    ]
    add_three_line_table(doc, ["项目", "经复核信息"], facts, [4.0, 9.5], 10.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(20)
    set_run_font(p.add_run("2026年8月"), 12, False, DESIGN["muted"])
    doc.add_page_break()


def build_document(output_path: Path):
    audit = json.loads((ARTIFACTS / "data_audit.json").read_text(encoding="utf-8"))
    refs = json.loads((ARTIFACTS / "verified_references.json").read_text(encoding="utf-8"))
    s = audit["snapshot"]

    doc = Document()
    configure_document(doc)
    cover_page(doc, audit)

    add_heading(doc, "摘  要", 1)
    add_body(doc, "传统行业代码能够稳定描述企业的登记主业，却难以反映多元经营企业已经开展的体育业务。本文以企业名称、行业代码和主要业务活动为基础，构建文本证据与行业代码协同的规则融合框架，并以业务线为分析单元识别体育业态及跨界经营。研究对象为四川省76,687条脱敏企业登记记录，原始数据包含统一社会信用代码、企业名称、行业代码和主要业务活动四个字段。", bold_prefix="传统行业代码")
    add_body(doc, "SportFusion先切分主要业务活动文本，再完成中文分词、体育词典匹配、行业代码类型判定与业态归类。项目以业务范围占比、关键词密度、行业代码权重和业态覆盖度四项特征构造SportRatio。该指标是0—1之间的相对业务比重代理值，不等同于经审计的体育营业收入占比。对企业SportRatio乘以100后汇总形成相对产出指数，指数只用于同一数据口径内的结构比较，不代表营业收入、增加值或货币产值。", bold_prefix="SportFusion")
    add_body(doc, f"正式重跑批次{s['formal_batch']}识别体育相关企业{s['fusion_sport_enterprises']:,}家，占全部样本的{s['fusion_coverage_pct']:.2f}%；传统直接行业代码法识别{s['traditional_sport_enterprises']:,}家，占{s['traditional_coverage_pct']:.2f}%。融合法比传统法多发现{s['incremental_enterprises']:,}家，识别数量相对增加{s['relative_identification_increase_pct']:.2f}%。体育企业中跨界经营企业{s['crossover_enterprises']:,}家，占{s['crossover_rate_among_sport_enterprises_pct']:.2f}%。项目汇总相对产出指数为{s['total_output_index']:,.2f}，体育企业的平均SportRatio为{s['average_sport_ratio_pct_among_sport_enterprises']:.2f}%。这些结果均由全量重跑文件直接复算。", bold_prefix="正式重跑批次")
    add_body(doc, "研究没有人工金标准、逐条预测分数和双人标注记录，因此不报告准确率、召回率、F1、AUC或Kappa。原稿中缺少支撑材料的监督评估指标、货币效益、趋势预测和吞吐性能数字均已删除或降为待验证事项。本文的贡献在于形成一套可解释的冷启动识别流程、可复算的数据证据链和可运行的统计辅助原型，并明确其适用边界。", bold_prefix="研究没有人工金标准")
    add_body(doc, "关键词：体育产业统计；文本识别；多元经营；业务边界；相对产出指数；证据审计", no_indent=True)
    add_callout(doc, "口径提示｜", "文中“产业规模”均指基于SportRatio汇总形成的相对规模指数；在未接入企业营收或增加值数据前，不作金额化解释。", DESIGN["warm_fill"])

    add_heading(doc, "目  录", 1)
    add_toc(doc)
    doc.add_page_break()

    # 第一章
    add_heading(doc, "第一章  研究问题与总体设计", 1)
    add_heading(doc, "1.1 研究背景与问题提出", 2)
    add_body(doc, "GB/T 4754—2017为国民经济活动提供统一分类基础，体育产业统计分类（2019）进一步界定体育产业统计范围。两类标准解决的是统一归类问题，但企业登记的行业代码通常对应主业。当企业同时经营赛事、培训、用品销售或健身服务时，单一代码无法完整呈现业务组合，统计边界由此出现遗漏。相关标准及政策依据均采用官方公开页面核验[1-4]。", bold_prefix="GB/T 4754—2017")
    add_body(doc, "文本字段为识别遗漏提供了补充证据。“主要业务活动”直接记录企业申报的经营内容，可拆分为若干业务线，再与体育业态词典和行业代码交叉判断。文本作为经济研究数据的处理原则强调测量过程、可重复性和外部有效性[5]；因此，本项目把可解释规则和证据留痕置于模型复杂度之前。", bold_prefix="文本字段")
    add_body(doc, "本研究关注三个可检验问题：其一，融合文本证据后，体育相关企业的识别范围改变多少；其二，在缺少分业务营收时，能否给出透明且可复算的相对业务权重；其三，如何将识别、汇总、风险提示和报告输出整合为可交付的统计辅助工具。")
    add_figure(doc, FIGURES / "01_研究论证闭环.png", "图1-1  研究问题、方法与验证闭环", "本研究根据正式批次和证据审计流程绘制。")

    add_heading(doc, "1.2 核心概念与统计边界", 2)
    add_table_caption(doc, "表1-1  核心概念与本文口径")
    add_three_line_table(doc, ["概念", "本文定义", "不作何种解释"], [
        ["体育相关企业", "达到融合规则阈值且具有可追溯体育文本或代码证据的企业", "不等同于法定统计名录认定"],
        ["SportRatio", "由四维特征加权形成的0—1相对代理值", "不等同于财务报表中的收入占比"],
        ["相对产出指数", "体育企业SportRatio×100后的汇总值", "不等同于营业收入、增加值或货币产值"],
        ["跨界经营", "行业代码并非直接体育代码，但文本证据显示存在体育业务", "不代表企业经营许可或合规结论"],
    ], [3.0, 7.0, 4.0])
    add_body(doc, "国家统计局发布的体育产业统计分类（2019）设置11个大类[2]。本项目没有改写官方分类，而是为规则识别和可视化汇总建立9个操作性业态标签。两者用途不同：官方分类服务统计制度，项目标签服务文本证据整理，成果应用时应通过映射表回接正式分类。", bold_prefix="国家统计局")

    add_heading(doc, "1.3 研究目标与创新点", 2)
    add_bullet(doc, "创新点一：", "形成零人工标签条件下可运行的双通道规则框架。行业代码提供结构先验，主要业务活动文本提供实际经营证据，输出同时保留命中词、业务线和判定依据。")
    add_bullet(doc, "创新点二：", "把企业级二值归类细化为业务线级特征测算。SportRatio保留经营内容的连续差异，并明确其代理变量属性。")
    add_bullet(doc, "创新点三：", "将“数据审计—全量重跑—冲突决策—图表引用”组织为可复现证据链。正文数字全部绑定正式批次，避免预处理记录、截断案例和最终结果混用。")
    add_bullet(doc, "创新点四：", "实现识别、产业分析、风险提示、证据追踪和结果导出的Web原型，为统计人员提供可视化操作入口。")
    add_body(doc, "上述创新首先是框架创新和实践创新。由于当前尚未完成金标准标注与外部财务校准，本文不把规则覆盖差异表述为算法准确率，也不把相对指数表述为真实经济总量。")

    # 第二章
    add_heading(doc, "第二章  数据基础与治理过程", 1)
    add_heading(doc, "2.1 数据来源与字段结构", 2)
    add_body(doc, "研究使用合作单位提供的脱敏企业登记信息。原始Excel含76,687条记录，核心字段为统一社会信用代码、详细名称、行业代码和主要业务活动。审计时同时记录文件大小、修改时间和SHA-256哈希，以保证后续复算能够定位到同一数据版本。")
    add_table_caption(doc, "表2-1  原始数据字段及研究用途")
    add_three_line_table(doc, ["字段", "用途", "处理原则"], [
        ["统一社会信用代码", "企业记录去重与结果回接", "仅作脱敏数据中的唯一标识"],
        ["详细名称", "名称词辅助、区域字符串抽取", "区域结果不得替代注册地址"],
        ["行业代码", "直接体育、间接相关和其他类型判定", "按GB/T 4754—2017规则映射"],
        ["主要业务活动", "业务线切分、分词、词典匹配和业态归类", "保留原文并生成处理字段"],
    ], [3.2, 5.0, 5.8])

    add_heading(doc, "2.2 数据质量审计", 2)
    q = audit["raw_data_quality"]
    add_table_caption(doc, "表2-2  原始数据质量检查结果")
    add_three_line_table(doc, ["检查项", "结果", "处理决定"], [
        ["记录总数", f"{q['rows']:,}条", "作为全量分析分母"],
        ["完全重复记录", f"{q['duplicate_rows']}条", "无须删除"],
        ["重复统一社会信用代码", f"{q['duplicate_credit_codes']}条", "唯一键检查通过"],
        ["主要业务活动缺失", f"{q['missing_business_activity']}条", "不进行文本推断，保留代码通道"],
        ["行业代码缺失", f"{q['missing_industry_code']}条", "本批次未发现"],
    ], [4.2, 3.0, 6.8])
    add_body(doc, f"主要业务活动缺失{q['missing_business_activity']}条，占全样本约{q['missing_business_activity']/q['rows']*100:.2f}%。这些记录无法提供文本证据，系统不以企业名称强行补造业务内容；若行业代码存在，则仅保留代码先验。该处理降低了误判风险，也使缺失值影响可被明确追踪。")
    add_figure(doc, FIGURES / "02_数据处理流程.png", "图2-1  数据治理与可追溯处理流程", "本研究依据原始Excel、处理后CSV和正式重跑文件绘制。")

    add_heading(doc, "2.3 文本处理与业态映射", 2)
    add_body(doc, "文本处理依次执行字符规范化、业务线切分、中文分词、体育关键词匹配和业态归类。中文分词由项目实际使用的jieba组件完成，并支持加载自定义词典[7]。词项权重设计参考经典文本检索中的词项加权思想[6]，但本项目没有训练统计分类器。")
    add_body(doc, "操作性业态包括体育用品、健身休闲、体育赛事、体育培训、体育场馆、体育管理、电子竞技、体育传媒和体育彩票。其作用是将命中证据归并为稳定的分析标签；正式统计报表仍需依据体育产业统计分类（2019）进行复核映射。")

    add_heading(doc, "2.4 数据适用边界", 2)
    add_body(doc, "登记文本反映企业申报经营内容，不直接等同于实际收入、就业或增加值。企业可能登记但尚未开展某项业务，也可能出现文本更新滞后。SportFusion适合用于名录扩展、线索发现和结构监测，不应替代统计调查、税务数据或企业财务报表。", bold_prefix="登记文本")
    add_body(doc, "区域字段由企业名称中的地名字符串抽取。该方法可以形成初步空间分布，但无法处理名称不含地名、名称地名与注册地址不一致以及省本级记录等情况。因此，区域结果只用于探索性比较，正式发布前必须回接登记地址字段。")

    # 第三章
    add_heading(doc, "第三章  双通道识别与比重测算方法", 1)
    add_heading(doc, "3.1 文本—代码双通道框架", 2)
    add_body(doc, "代码通道回答“登记主业是否与体育直接相关”，文本通道回答“企业实际申报了哪些体育业务”。两条通道先独立形成证据，再由规则层进行融合。这样的设计可以定位每一次判定来自何种代码、何条业务线和哪些关键词，便于统计人员复核。")
    add_figure(doc, FIGURES / "03_双通道识别框架.png", "图3-1  文本与行业代码双通道识别框架", "本研究依据backend/services/sport_recognition.py的实际逻辑绘制。")

    add_heading(doc, "3.2 业务线识别与最终判定", 2)
    add_body(doc, "系统将主要业务活动按分隔符切分为业务线，并逐条匹配体育词典。代码类型分为直接体育、间接相关和其他三类；体育相关判定要求SportRatio不低于0.10且形成主要体育业态，或由直接体育代码提供高强度先验。阈值是当前工程规则，不是通过人工金标准优化得到的最优阈值。")
    add_table_caption(doc, "表3-1  双通道证据组合及解释")
    add_three_line_table(doc, ["代码证据", "文本证据", "处理", "解释"], [
        ["直接体育", "命中", "识别为体育相关", "代码与经营文本相互支持"],
        ["直接体育", "未命中", "保留代码通道并标记复核", "可能存在文本缺失或登记滞后"],
        ["间接相关", "命中", "识别并标记跨界", "文本补足代码边界"],
        ["其他", "命中", "达到阈值后识别并标记跨界", "由文本发现的非主业体育经营"],
        ["其他", "未命中", "不识别", "缺少体育业务证据"],
    ], [2.4, 2.4, 4.4, 5.0])

    add_heading(doc, "3.3 SportRatio四维加权", 2)
    add_body(doc, "在缺少分业务营收数据时，项目以四项可解释特征构造相对代理值。W1表示体育业务线占全部业务线的比例；W2表示归一化后的体育关键词密度；W3表示行业代码先验，直接体育、间接相关和其他分别取0.85、0.30和0；W4表示命中的体育业态数占项目九类业态的比例。")
    add_formula(doc, "SportRatio = 0.40W1 + 0.25W2 + 0.25W3 + 0.10W4")
    add_table_caption(doc, "表3-2  SportRatio特征定义")
    add_three_line_table(doc, ["特征", "计算口径", "组合权重", "证据含义"], [
        ["W1 业务范围", "体育业务线数/业务线总数", "0.40", "经营范围中的体育内容覆盖"],
        ["W2 关键词密度", "min(体育词命中数/分词数×10, 1)", "0.25", "体育语义在文本中的集中程度"],
        ["W3 代码权重", "直接0.85；间接0.30；其他0", "0.25", "现行行业分类的结构先验"],
        ["W4 业态覆盖", "命中体育业态数/9", "0.10", "体育经营内容的多样程度"],
    ], [3.0, 5.2, 2.2, 4.0])
    add_figure(doc, FIGURES / "04_SportRatio测算流程.png", "图3-2  SportRatio四维加权测算流程", "本研究依据正式运行代码绘制；权重为人工设定规则权重。")
    add_callout(doc, "解释边界｜", "SportRatio只在当前规则、词典和数据字段下具有相对比较意义。若取得企业分业务营收，应优先以真实营收占比校准或替代该代理值。")

    add_heading(doc, "3.4 相对产出指数与汇总指标", 2)
    add_body(doc, "对每家体育相关企业，将SportRatio乘以100得到企业相对产出指数，再按业态或区域求和。总指数579,124.95等于所有体育企业SportRatio之和乘以100。由于每家企业的基准规模均被规范化为100，该指数没有货币单位，不能据此推导产业产值。")
    add_formula(doc, "企业相对产出指数_i = SportRatio_i × 100")
    add_formula(doc, "分组相对产出指数_g = Σ(企业相对产出指数_i)，i∈g")
    add_body(doc, "CR3、CR5、HHI和Gini用于描述相对指数在区域间的集中程度。它们是当前名称抽取口径下的结构指标，不能直接替代基于注册地址和真实增加值计算的区域产业集中度。")

    # 第四章
    add_heading(doc, "第四章  可复现验证与证据审计", 1)
    add_heading(doc, "4.1 验证原则", 2)
    add_body(doc, "本研究采用“复算优先、交叉核对、无证据降级”的验证原则。A类证据来自原始数据直接统计，B类证据来自处理后明细或运行代码输出，C类证据来自可核验的政府标准与原始文献，D类为来源不明、批次冲突或无法复算的数字。A—C类可进入结论，D类必须删除、重算或明确降级。")
    add_figure(doc, FIGURES / "10_证据分级与处理规则.png", "图4-1  数字证据分级与处理规则", "本研究数据审计规则。图中D级数字仅为被删除的旧稿示例，不构成本研究结果。")

    add_heading(doc, "4.2 正式批次与口径锁定", 2)
    add_body(doc, f"正式结果来自批次{s['formal_batch']}。识别脚本在76,687条全量处理数据上重新运行，随后由分析脚本生成业态、区域和集中度结果。审计程序再对总数、分组加总和文件哈希执行自动检查；业态企业数加总为{s['fusion_sport_enterprises']:,}，业态相对产出指数加总为{s['total_output_index']:,.2f}，均与总览一致。")
    add_table_caption(doc, "表4-1  旧稿冲突与修订决策")
    add_three_line_table(doc, ["冲突事项", "统一口径", "修订理由"], [
        ["体育企业总数", "8,950家", "9,023家属于预处理初筛文件，正式全量识别结果为8,950家"],
        ["增量企业数", "934家", "100仅为JSON中截断案例列表长度，不是总体增量"],
        ["11.7%的含义", "识别数量相对增加11.65%", "没有人工金标准，不能写为准确率提升"],
        ["监督评估指标", "不报告", "缺少逐条人工标签、预测分数和标注者记录"],
        ["579,124.95的含义", "相对产出指数", "是SportRatio×100后的汇总，不是货币金额"],
        ["趋势与经济效益", "降为待验证场景", "缺少多期数据、成本台账与对照试验"],
    ], [3.4, 4.0, 6.8])

    add_heading(doc, "4.3 基线比较的有效解释", 2)
    add_body(doc, "传统直接行业代码法与融合识别法在同一全量样本上比较，能够回答识别范围是否扩展，却不能回答新增企业中有多少为真实阳性。换言之，934家是规则意义上的增量发现，不是经人工核验的漏检纠正数。只有建立分层抽样金标准后，才可报告准确率、精确率、召回率、F1等指标。", bold_prefix="传统直接行业代码法")
    add_body(doc, "当前验证包括四类内部一致性检查：原始行数与处理后行数一致；业态分组与总量一致；传统法与融合法使用同一分母；图表、表格和摘要绑定同一审计JSON。这些检查保障结果可复算，但不替代外部效度检验。")

    add_heading(doc, "4.4 可复现文件链", 2)
    add_table_caption(doc, "表4-2  关键证据文件与用途")
    add_three_line_table(doc, ["文件", "用途", "SHA-256前16位"], [
        ["企业原始数据.xlsx", "原始76,687条企业登记记录", audit["sources"][0]["sha256"][:16]],
        ["enterprise_dataset_20260629_160902.csv", "分词、关键词及代码类型处理结果", audit["sources"][1]["sha256"][:16]],
        ["sport_ratio_results_20260801_203307.csv", "正式全量识别与比重结果", audit["sources"][3]["sha256"][:16]],
        ["industry_analysis_20260801_203329.json", "业态、区域与集中度汇总", audit["sources"][5]["sha256"][:16]],
        ["data_audit.json", "正文统一口径与冲突决策", "由上述文件交叉生成"],
    ], [6.5, 5.0, 3.0])

    # 第五章
    add_heading(doc, "第五章  实证结果与结构分析", 1)
    add_heading(doc, "5.1 识别范围变化", 2)
    add_body(doc, f"传统直接行业代码法识别{s['traditional_sport_enterprises']:,}家，融合识别法识别{s['fusion_sport_enterprises']:,}家。两者分别占全样本的{s['traditional_coverage_pct']:.2f}%和{s['fusion_coverage_pct']:.2f}%，差额为{s['incremental_enterprises']:,}家，相当于传统识别数量的{s['relative_identification_increase_pct']:.2f}%。这一结果说明文本字段能补充行业代码未呈现的经营内容。")
    add_table_caption(doc, "表5-1  两种识别方法的范围比较")
    add_three_line_table(doc, ["方法", "识别企业数", "占全样本", "相对传统法"], [
        ["传统直接行业代码法", f"{s['traditional_sport_enterprises']:,}", f"{s['traditional_coverage_pct']:.2f}%", "基线"],
        ["文本—代码融合识别法", f"{s['fusion_sport_enterprises']:,}", f"{s['fusion_coverage_pct']:.2f}%", f"+{s['incremental_enterprises']:,}家（+{s['relative_identification_increase_pct']:.2f}%）"],
    ], [4.5, 3.0, 3.0, 4.2])
    add_figure(doc, FIGURES / "05_识别范围对比.png", "图5-1  传统行业代码法与融合识别法的识别范围对比", f"正式批次{s['formal_batch']}；本研究计算。")

    add_heading(doc, "5.2 体育业务相对比重分布", 2)
    dist = audit["ratio_distribution"]
    add_body(doc, f"全样本中SportRatio为0的企业有{dist['0']:,}家；(0,0.2]区间有{dist['0-0.2']:,}家。低比重记录被完整保留，说明系统不是先删去非体育企业再画分布。体育相关企业的平均SportRatio为{s['average_sport_ratio_pct_among_sport_enterprises']:.2f}%，该均值不适用于全部76,687家企业。")
    add_figure(doc, FIGURES / "06_体育业务占比分布.png", "图5-2  全样本SportRatio区间分布", f"正式批次{s['formal_batch']}；本研究计算。")
    add_table_caption(doc, "表5-2  SportRatio分布明细")
    add_three_line_table(doc, ["区间", "企业数", "占全样本"], [
        ["0", f"{dist['0']:,}", f"{dist['0']/s['total_enterprises']*100:.2f}%"],
        ["(0, 0.2]", f"{dist['0-0.2']:,}", f"{dist['0-0.2']/s['total_enterprises']*100:.2f}%"],
        ["(0.2, 0.5]", f"{dist['0.2-0.5']:,}", f"{dist['0.2-0.5']/s['total_enterprises']*100:.2f}%"],
        ["(0.5, 0.8]", f"{dist['0.5-0.8']:,}", f"{dist['0.5-0.8']/s['total_enterprises']*100:.2f}%"],
        ["(0.8, 1.0]", f"{dist['0.8-1.0']:,}", f"{dist['0.8-1.0']/s['total_enterprises']*100:.2f}%"],
    ], [5.0, 4.0, 4.0])

    add_heading(doc, "5.3 业态结构与跨界经营", 2)
    add_body(doc, f"9个操作性业态的企业数加总为{s['fusion_sport_enterprises']:,}家，相对产出指数加总为{s['total_output_index']:,.2f}。体育用品与健身休闲合计占相对产出指数的{audit['category_distribution'][0]['output_share_pct'] + audit['category_distribution'][1]['output_share_pct']:.2f}%，构成当前样本的主体。电子竞技、体育传媒和体育彩票样本量较小，其高跨界率易受小样本影响，不宜直接外推。")
    add_figure(doc, FIGURES / "07_业态结构与跨界率.png", "图5-3  体育业态相对产出结构与跨界经营率", f"正式批次{s['formal_batch']}；相对产出指数为非货币单位。")
    add_table_caption(doc, "表5-3  各操作性业态的识别与相对产出结果")
    category_rows = []
    for item in audit["category_distribution"]:
        category_rows.append([
            item["category"], f"{item['enterprise_count']:,}", f"{item['output_index']:,.2f}",
            f"{item['output_share_pct']:.2f}%", f"{item['avg_sport_ratio']*100:.2f}%",
            f"{item['crossover_count']:,}", f"{item['crossover_pct']:.2f}%",
        ])
    add_three_line_table(doc, ["业态", "企业数", "相对产出指数", "指数份额", "平均比重", "跨界数", "跨界率"], category_rows, [2.3, 1.8, 2.8, 2.1, 2.1, 1.8, 1.9], 8.4)
    add_body(doc, f"体育赛事的跨界率为36.10%，高于体育用品、健身休闲、体育培训和体育场馆，说明赛事业务更常嵌入文旅、会展或综合服务企业。体育企业中共有{s['crossover_enterprises']:,}家被标记为跨界，占{s['crossover_rate_among_sport_enterprises_pct']:.2f}%。这一比例来自规则定义，后续仍需人工抽样核验。")

    add_heading(doc, "5.4 区域分布与集中度", 2)
    add_body(doc, f"按企业名称字符串抽取的地级市结果中，成都市相对产出指数为{audit['city_top10'][0]['产出指数']:,.2f}，体育企业数为{audit['city_top10'][0]['体育企业数']:,}家。全部区域口径包含“四川省本级”，据此计算CR3为{s['spatial_concentration_all_regions']['cr3_pct']:.2f}%，CR5为{s['spatial_concentration_all_regions']['cr5_pct']:.2f}%，HHI为{s['spatial_concentration_all_regions']['hhi']:.2f}，Gini为{s['spatial_concentration_all_regions']['gini']:.6f}。")
    add_figure(doc, FIGURES / "08_区域相对产出指数.png", "图5-4  地级市体育业务相对产出指数前十位", f"正式批次{s['formal_batch']}；区域由企业名称抽取，四川省本级未列入地级市排名。")
    add_table_caption(doc, "表5-4  地级市相对产出指数前十位")
    city_rows = []
    for idx, item in enumerate(audit["city_top10"], 1):
        city_rows.append([idx, item["区域"], f"{item['体育企业数']:,}", f"{item['产出指数']:,.2f}", f"{item['平均体育占比']*100:.2f}%", item["主导业态"]])
    add_three_line_table(doc, ["序号", "区域", "体育企业数", "相对产出指数", "平均比重", "主导业态"], city_rows, [1.2, 2.4, 2.4, 3.0, 2.2, 3.0], 8.8)
    add_callout(doc, "区域解释边界｜", "名称抽取不能替代注册地址。正式区域统计应以登记机关或标准地址编码回接；本结果只说明当前字符串规则下的结构特征。")

    add_heading(doc, "5.5 结果讨论", 2)
    add_body(doc, "融合识别的主要价值在于发现传统代码口径以外的候选企业，并给出可复核证据。934家增量记录使统计人员能够建立重点复核清单，而非直接改变法定统计总量。")
    add_body(doc, "SportRatio把二值归类扩展为连续代理变量，使同一业态内部的经营内容差异能够进入汇总。但代理值仍由登记文本和人工权重共同决定，不能据此判断企业真实体育营收。")
    add_body(doc, "区域和新兴业态结果显示了进一步补充数据的方向：区域分析需要标准地址，新兴业态需要扩大样本和建立人工标注，规模测算需要接入营收或增加值。结果的价值不仅在于给出数字，也在于明确下一步数据治理的优先级。")

    # 第六章
    add_heading(doc, "第六章  系统架构与原型实现", 1)
    add_heading(doc, "6.1 技术架构", 2)
    add_body(doc, "系统采用前后端分离架构。交互层使用Vue 3与Element Plus，接口层由FastAPI提供数据预处理、企业识别、比重测算、产业分析和风险快照等服务，服务层封装NLP预处理与规则融合逻辑，数据层保留原始Excel、处理后CSV/JSON、SQLite及批次信息。")
    add_figure(doc, FIGURES / "09_系统技术架构.png", "图6-1  SportFusion系统技术架构", "依据项目当前前后端代码与路由结构绘制。")
    add_table_caption(doc, "表6-1  系统主要模块与统计用途")
    add_three_line_table(doc, ["模块", "核心功能", "统计用途"], [
        ["企业数据治理", "上传、清洗、字段校验与批次管理", "建立可追溯数据底表"],
        ["企业业务识别", "单条与批量文本—代码融合判定", "形成候选企业及判定依据"],
        ["经营比重测算", "四维特征计算与企业级解释", "生成相对业务比重代理值"],
        ["产业规模分析", "业态、区域、集中度与漏斗分析", "支持结构比较与复核排序"],
        ["风险与证据追踪", "批次异常、口径差异和来源展示", "降低误用与版本混淆"],
        ["报告与成果中心", "表格、图表和结果导出", "服务复核、汇报与留档"],
    ], [3.5, 6.0, 5.0])

    add_heading(doc, "6.2 真实运行界面", 2)
    add_body(doc, "以下截图由本项目当前前后端在本地启动后自动抓取，能够证明原型界面和主要功能入口可运行。页面保留“真实数据优先、演示保留可用”的提示，因此截图只用于说明系统实现，不作为独立的数据证据。正文实证数字仍以正式重跑文件为准。", bold_prefix="以下截图")
    add_figure(doc, SCREENSHOTS / "01_监测驾驶舱.png", "图6-2  统计监测驾驶舱运行界面", "2026年8月1日本地运行截图；页面数据模式以界面标识为准。", 15.4)
    add_figure(doc, SCREENSHOTS / "02_企业识别.png", "图6-3  企业单条与批量识别运行界面", "2026年8月1日本地运行截图。", 15.4)
    add_figure(doc, SCREENSHOTS / "04_模型评估.png", "图6-4  模型评估页面的证据约束状态", "2026年8月1日本地运行截图；当前正式批次未生成监督评估结果，因此页面不展示虚构指标。", 15.4)

    add_heading(doc, "6.3 统计业务落地流程", 2)
    add_body(doc, "建议将系统嵌入“机器初筛—业务复核—名录回接—定期更新”的统计流程。机器初筛负责生成候选企业、业务线证据和相对比重；业务人员对边界样本和高影响企业复核；确认结果回接统计名录；词典、规则和批次记录随新业态持续更新。")
    add_body(doc, "对外发布前应设置三道闸门：数据版本冻结，人工抽样复核，指标含义审签。任何金额化结果都必须接入真实营收或增加值，并记录价格口径、核算期间和缺失值处理。")

    # 第七章
    add_heading(doc, "第七章  应用价值、政策建议与局限", 1)
    add_heading(doc, "7.1 可确认的应用价值", 2)
    add_body(doc, "第一，系统能够将文本证据转化为可复核的候选名录。本批次相对传统代码法新增934家体育相关候选企业，其中跨界标记可帮助统计人员集中核验高遗漏风险对象。")
    add_body(doc, "第二，系统能够统一保存企业级证据和批次信息。相比只留汇总表，明细化证据链更适合处理规则调整、企业申诉和历史追溯。")
    add_body(doc, "第三，系统能够在缺少财务数据时提供相对结构视图。业态份额、跨界率和区域集中度可以支持抽样设计、名录维护和调研资源排序，但不能代替官方总量核算。")

    add_heading(doc, "7.2 面向不同主体的使用方式", 2)
    add_table_caption(doc, "表7-1  多主体应用场景与边界")
    add_three_line_table(doc, ["主体", "可直接使用", "必须补充"], [
        ["统计部门", "候选名录、证据明细、抽样复核排序", "人工金标准、正式分类映射和财务核算数据"],
        ["体育主管部门", "业态结构、跨界线索和区域调研清单", "企业实际经营核验与政策对象资格审核"],
        ["园区或协会", "产业链主体发现与会员服务线索", "企业授权、动态经营信息和合规审查"],
        ["研究机构", "规则基线、复现实验和误差分析框架", "多期数据、外部样本与标注数据"],
        ["企业", "自身业务描述规范化和同类比较参考", "不得将系统分数用于信用或行政认定"],
    ], [2.8, 6.0, 6.0])

    add_heading(doc, "7.3 政策与实施建议", 2)
    add_bullet(doc, "建立分层复核样本。", "按代码类型、SportRatio区间、业态和区域分层抽取样本，至少由两名业务人员独立标注，并保存分歧处置记录。")
    add_bullet(doc, "建设动态词典治理机制。", "新增词条应记录来源、适用业态、启用时间和退役原因，避免词典扩张造成不可解释的口径漂移。")
    add_bullet(doc, "回接标准地址与正式分类。", "区域统计使用注册地址或标准地理编码；操作性九业态通过映射表对接体育产业统计分类（2019）。")
    add_bullet(doc, "以真实财务数据校准规模。", "在企业授权和数据安全框架内接入分业务营收、增加值或税务数据，检验SportRatio与真实业务权重的关系。")
    add_bullet(doc, "固化版本与审签制度。", "每次发布同步保存数据哈希、代码版本、参数、规则词典和审签人，确保同一结果可以重算。")

    add_heading(doc, "7.4 局限性", 2)
    add_body(doc, "数据局限：样本为单一时点的企业登记信息，不能形成可靠趋势；主要业务活动可能滞后于真实经营；区域由名称字符串抽取。")
    add_body(doc, "方法局限：当前系统采用词典和规则融合，不能充分理解隐喻、否定或复杂上下文；四项组合权重由工程经验设定，尚无财务真值校准。")
    add_body(doc, "验证局限：工作区缺少逐条人工金标准、标注者一致性记录和外部测试集。因此，本文只能验证数据一致性、规则覆盖变化和内部可复算性，不能证明统计分类准确率。")
    add_body(doc, "应用局限：相对产出指数不是货币指标，不能直接进入GDP、产业增加值或经济效益核算。系统原型保留演示回退，正式部署应关闭演示数据并启用权限、日志和审签机制。")

    # 第八章
    add_heading(doc, "第八章  结论与展望", 1)
    add_heading(doc, "8.1 主要结论", 2)
    add_body(doc, f"第一，文本—代码融合规则在同一76,687家企业样本上识别{s['fusion_sport_enterprises']:,}家体育相关企业，比传统直接行业代码法增加{s['incremental_enterprises']:,}家，识别数量相对增加{s['relative_identification_increase_pct']:.2f}%。该结果证明文本字段能够拓展候选识别范围，但不等同于准确率提升。")
    add_body(doc, f"第二，SportRatio为多元经营企业提供连续的相对业务比重代理值。体育企业平均值为{s['average_sport_ratio_pct_among_sport_enterprises']:.2f}%，汇总相对产出指数为{s['total_output_index']:,.2f}。二者均基于登记文本与规则权重，不能解释为真实营收占比或货币产业规模。")
    add_body(doc, f"第三，项目识别跨界经营企业{s['crossover_enterprises']:,}家，占体育相关企业的{s['crossover_rate_among_sport_enterprises_pct']:.2f}%。体育赛事等业态体现较强跨界特征；电子竞技和体育传媒跨界率较高，但样本量小，结论应谨慎。")
    add_body(doc, "第四，可复算证据链与可运行原型共同提升了成果的可交付性。数据哈希、正式批次、冲突处理和结果边界均进入报告，使统计使用者能够区分事实、规则输出与待验证假设。")

    add_heading(doc, "8.2 后续工作", 2)
    add_body(doc, "下一阶段的重点不是直接替换规则框架，而是建立可靠的校准基础。首先完成分层人工标注和双人一致性检验，再比较规则基线、传统机器学习和预训练语言模型；只有在同一外部测试集上，才能讨论模型优劣。")
    add_body(doc, "取得企业分业务营收或增加值后，可用真实业务权重校准SportRatio，并分别报告代理值误差和总量核算误差。建立多年度同口径样本后，才适合开展趋势分析和预测。")
    add_body(doc, "系统部署方面，应增加权限分级、操作日志、词典版本管理、数据脱敏和异常回滚；对外报告自动生成时，应强制附带批次号、数据来源、指标定义和适用边界。")

    # 参考文献
    add_heading(doc, "参考文献", 1)
    ref_order = ["R1", "R2", "R3", "R4", "R5", "R6", "R7"]
    ref_map = {r["id"]: r for r in refs}
    for idx, rid in enumerate(ref_order, 1):
        r = ref_map[rid]
        author = r.get("authors") or r.get("publisher")
        identifier = r.get("identifier", "")
        text = f"[{idx}] {author}. {r['title']}[{('EB/OL' if r['url'].startswith('http') else 'M')}]. {r['publisher']}, {r['year']}. {identifier}. {r['url']}"
        p = add_body(doc, text, no_indent=True)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.space_after = Pt(5)

    # 附录
    add_heading(doc, "附录A  指标定义与计算口径", 1)
    add_table_caption(doc, "表A-1  核心指标定义")
    add_three_line_table(doc, ["指标", "公式或定义", "使用限制"], [
        ["融合覆盖率", "融合识别企业数/全样本企业数", "描述规则识别范围，不是召回率"],
        ["相对识别增加", "(融合法数量−传统法数量)/传统法数量", "不代表准确率提升"],
        ["跨界率", "跨界企业数/体育相关企业数", "依赖当前代码和文本规则"],
        ["平均SportRatio", "体育企业SportRatio均值", "只针对体育企业子集"],
        ["相对产出指数", "Σ(SportRatio×100)", "非货币单位"],
        ["CR3/CR5", "前3/前5区域指数占全部区域指数", "区域抽取需地址数据校准"],
        ["HHI", "各区域指数份额平方和×10,000", "反映当前相对指数集中度"],
        ["Gini", "区域指数分布的基尼系数", "受省本级与地级市混合口径影响"],
    ], [3.4, 6.0, 5.0])

    add_heading(doc, "附录B  复现步骤", 1)
    add_body(doc, "1. 使用enterprise_dataset_20260629_160902.csv运行backend/scripts/run_recognition.py，输出全量SportRatio结果与模型比较JSON。", no_indent=True)
    add_body(doc, "2. 使用正式SportRatio结果运行backend/scripts/run_analysis.py，输出业态、区域、集中度和汇总JSON。", no_indent=True)
    add_body(doc, "3. 运行paper_revision/run_data_audit.py，对原始Excel、正式结果和历史输出进行行数、加总、哈希及冲突检查。", no_indent=True)
    add_body(doc, "4. 图表统一读取paper_revision/artifacts/data_audit.json；正文不直接抄录历史文档中的孤立数字。", no_indent=True)
    add_body(doc, "5. 对外发布前补充人工标注与财务校准；未完成前，不生成监督评估指标和金额化结论。", no_indent=True)

    add_heading(doc, "附录C  数据真实性与修订声明", 1)
    add_body(doc, "本修订稿仅保留能够由工作区原始数据、正式重跑结果、运行代码或官方公开来源直接核验的事实。旧稿中的冲突数字按正式批次统一；无法定位原始标签、测试记录、财务台账或时间序列的数据均不作为实证结论。")
    add_body(doc, "报告保留研究局限，不以界面展示替代数据证据，不以规则输出替代人工金标准，不以相对指数替代货币规模。读者可依据附录B和文件哈希复核主要结果。")

    # Final sweep: lock all runs to Songti / Times New Roman and validate text.
    all_text = []
    for paragraph in doc.paragraphs:
        all_text.append(paragraph.text)
        for run in paragraph.runs:
            set_run_font(run, run.font.size.pt if run.font.size else None, run.bold, None)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text.append(paragraph.text)
                    for run in paragraph.runs:
                        set_run_font(run, run.font.size.pt if run.font.size else None, run.bold, None)
    assert_no_unsupported_assertions("\n".join(all_text))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = REVISION / "体融识界·SportFusion_国赛优化稿.docx"
    path = build_document(out)
    print(json.dumps({"output": str(path), "size_bytes": path.stat().st_size}, ensure_ascii=False))
