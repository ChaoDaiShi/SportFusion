from __future__ import annotations

import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "paper_revision" / "体融识界·SportFusion_国赛优化稿.docx"


def iter_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def validate():
    doc = Document(DOCX)
    errors = []
    warnings = []
    paragraphs = list(iter_paragraphs(doc))
    text = "\n".join(p.text for p in paragraphs)

    # Font audit: all visible text runs must carry the requested font pair.
    checked_runs = 0
    for p in paragraphs:
        for run in p.runs:
            if not run.text.strip():
                continue
            checked_runs += 1
            rpr = run._element.rPr
            rfonts = rpr.rFonts if rpr is not None else None
            east = rfonts.get(qn("w:eastAsia")) if rfonts is not None else None
            ascii_font = rfonts.get(qn("w:ascii")) if rfonts is not None else None
            if east != "宋体":
                errors.append(f"中文字体未锁定：{run.text[:24]} -> {east}")
            if ascii_font != "Times New Roman":
                errors.append(f"西文字体未锁定：{run.text[:24]} -> {ascii_font}")

    # Tables with multiple rows are data tables and must be three-line tables.
    data_tables = 0
    for table in doc.tables:
        if len(table.rows) <= 1:
            continue
        data_tables += 1
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is None:
            errors.append("数据表缺少表级边框设置")
            continue
        names = {child.tag.split("}")[-1] for child in borders}
        if not {"top", "bottom"}.issubset(names):
            errors.append(f"数据表缺少顶线或底线：{names}")
        if names.intersection({"left", "right", "insideV", "start", "end"}):
            errors.append(f"数据表出现竖线：{names}")

    # Headings and caption numbering.
    headings = [p.text.strip() for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    duplicate_headings = [x for x, c in Counter(headings).items() if c > 1]
    if duplicate_headings:
        errors.append(f"标题重复：{duplicate_headings}")

    figure_captions = [p.text.strip() for p in doc.paragraphs if re.match(r"^图\d+-\d+", p.text.strip())]
    table_captions = [p.text.strip() for p in doc.paragraphs if re.match(r"^表(?:\d+|[A-Z])-\d+", p.text.strip())]
    for label, captions in (("图", figure_captions), ("表", table_captions)):
        ids = [re.match(rf"^{label}([^ ]+)", cap).group(1) for cap in captions]
        dup = [x for x, c in Counter(ids).items() if c > 1]
        if dup:
            errors.append(f"{label}编号重复：{dup}")

    # Truthfulness and terminology audit.
    required_strings = [
        "76,687", "8,950", "8,016", "934", "11.65%", "579,124.95",
        "相对产出指数", "非货币", "人工金标准", "正式批次20260801_203307",
    ]
    for item in required_strings:
        if item not in text:
            errors.append(f"缺少核心口径：{item}")
    banned_patterns = [
        r"AUC\s*[=:：]\s*0\.91",
        r"Kappa\s*[=:：]\s*0\.86",
        r"Pearson\s*r\s*[=:：]\s*0\.72",
        r"纠正约?6亿元",
        r"准确率提升11\.7%",
    ]
    for pattern in banned_patterns:
        if re.search(pattern, text, re.I):
            errors.append(f"出现未获支持的旧稿断言：{pattern}")

    # Embedded media audit.
    media_info = []
    with zipfile.ZipFile(DOCX) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        for name in media:
            data = zf.read(name)
            temp = ROOT / "paper_revision" / "artifacts" / "_media_probe.png"
            temp.write_bytes(data)
            try:
                with Image.open(temp) as im:
                    media_info.append({"name": name, "width": im.width, "height": im.height})
                    if im.width < 1400:
                        warnings.append(f"嵌图宽度偏小：{name} {im.width}x{im.height}")
            except Exception:
                warnings.append(f"无法读取嵌图尺寸：{name}")
            finally:
                temp.unlink(missing_ok=True)
    if len(media_info) < 13:
        errors.append(f"嵌图数量不足：{len(media_info)}")

    report = {
        "docx": str(DOCX),
        "size_bytes": DOCX.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "headings": len(headings),
        "tables_total": len(doc.tables),
        "data_tables_three_line_checked": data_tables,
        "figure_captions": len(figure_captions),
        "table_captions": len(table_captions),
        "inline_shapes": len(doc.inline_shapes),
        "embedded_media": len(media_info),
        "font_runs_checked": checked_runs,
        "errors": errors[:50],
        "warnings": warnings[:50],
        "status": "PASS" if not errors else "FAIL",
    }
    out = ROOT / "paper_revision" / "artifacts" / "final_docx_validation.json"
    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if errors:
        raise SystemExit(1)


if __name__ == "__main__":
    validate()
