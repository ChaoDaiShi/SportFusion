from __future__ import annotations

import csv
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "体融识界·SportFusion——基于NLP文本识别与多维度加权的多元经营企业体育业务边界识别与产业规模测算.docx"
ARTIFACTS = ROOT / "paper_revision" / "artifacts"

NUMBER_PATTERN = re.compile(r"(?<![A-Za-z])(?:\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)(?:%|％)?")
RISK_PATTERNS = {
    "evaluation": re.compile(r"准确率|精确率|召回率|F1|AUC|ROC|Kappa|Pearson|混淆矩阵|人工标注|n\s*=", re.I),
    "forecast": re.compile(r"预测|预计|趋势|202[0-9][—\-至到]202[0-9]|复合增长|增长率"),
    "economic_benefit": re.compile(r"节省|经济效益|万元|亿元|投入产出|资源错配|人工成本"),
    "performance": re.compile(r"响应时间|吞吐|并发|内存|CPU|压力测试|秒/万条|条/秒"),
    "concentration": re.compile(r"CR3|CR5|HHI|基尼", re.I),
    "output_value": re.compile(r"产值|增加值|营业收入|产出指数|产业规模"),
}


def extract_numeric_tokens(text: str) -> list[str]:
    return NUMBER_PATTERN.findall(text or "")


def risk_tags(text: str) -> list[str]:
    return [name for name, pattern in RISK_PATTERNS.items() if pattern.search(text or "")]


def decision_for(text: str, tags: list[str], tokens: list[str], formal_tokens: set[str]) -> tuple[str, str, str]:
    if "evaluation" in tags:
        return (
            "remove",
            "D",
            "当前工作区未提供逐条人工金标准、预测分数或标注者记录，不报告准确率、AUC、Kappa、Pearson相关系数等指标。",
        )
    if "forecast" in tags:
        return (
            "rewrite_as_scope",
            "D",
            "现有数据为单一时间截面，仅能描述当前样本结构，不能据此作确定性趋势预测。",
        )
    if "economic_benefit" in tags:
        return (
            "rewrite_as_scope",
            "D",
            "系统具备减少重复统计作业的潜力，实际经济效益需结合部署范围、人员工时和成本台账另行评估。",
        )
    if "performance" in tags:
        return (
            "remove",
            "D",
            "当前材料未提供可复核的压力测试原始记录，因此不报告平台性能数值。",
        )
    normalized = {token.replace(",", "").replace("％", "%") for token in tokens}
    if normalized and normalized.issubset(formal_tokens):
        return "keep", "A/B", "与2026-08-01正式复算快照一致。"
    return "recalculate", "D", "仅在与正式复算快照或经核验的权威来源一致时保留。"


def iter_claim_texts(document: Document):
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            yield f"paragraph:{index}", paragraph.style.name, text
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            text = " | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells)
            if text.strip(" |"):
                yield f"table:{table_index}:row:{row_index}", "Table Row", text


def main() -> None:
    ARTIFACTS.mkdir(parents=True, exist_ok=True)
    audit = json.loads((ARTIFACTS / "data_audit.json").read_text(encoding="utf-8"))
    snapshot = audit["snapshot"]
    formal_values = {
        snapshot["total_enterprises"],
        snapshot["traditional_sport_enterprises"],
        snapshot["fusion_sport_enterprises"],
        snapshot["crossover_enterprises"],
        snapshot["incremental_enterprises"],
        snapshot["traditional_coverage_pct"],
        snapshot["fusion_coverage_pct"],
        snapshot["relative_identification_increase_pct"],
        snapshot["average_sport_ratio_pct_among_sport_enterprises"],
        snapshot["total_output_index"],
        snapshot["crossover_rate_among_sport_enterprises_pct"],
        snapshot["spatial_concentration_all_regions"]["cr3_pct"],
        snapshot["spatial_concentration_all_regions"]["cr5_pct"],
        snapshot["spatial_concentration_all_regions"]["hhi"],
        snapshot["spatial_concentration_all_regions"]["gini"],
    }
    formal_tokens = set()
    for value in formal_values:
        if isinstance(value, int):
            formal_tokens.add(str(value))
        else:
            formal_tokens.update({str(value), f"{value:.1f}", f"{value:.2f}"})
            formal_tokens.update({f"{value}%", f"{value:.1f}%", f"{value:.2f}%"})

    document = Document(SOURCE_DOCX)
    rows = []
    headings = []
    captions = []
    for location, style, text in iter_claim_texts(document):
        if style.startswith("Heading"):
            headings.append((location, style, text))
        if re.match(r"^[图表]\s*\d+[\-—]\d+", text):
            captions.append((location, text))
        tokens = extract_numeric_tokens(text)
        if not tokens:
            continue
        tags = risk_tags(text)
        decision, grade, replacement = decision_for(text, tags, tokens, formal_tokens)
        rows.append(
            {
                "location": location,
                "style": style,
                "claim_text": text,
                "numeric_tokens": ";".join(tokens),
                "risk_tags": ";".join(tags),
                "evidence_grade": grade,
                "decision": decision,
                "replacement": replacement,
            }
        )

    fieldnames = list(rows[0].keys())
    for filename in ("docx_claims.csv", "claim_decisions.csv"):
        with (ARTIFACTS / filename).open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)

    heading_numbers = []
    for location, style, text in headings:
        match = re.match(r"^(\d+(?:\.\d+){0,2})\s*", text)
        if match:
            heading_numbers.append((match.group(1), location, text))
    duplicate_headings = {
        number: [text for item_number, _, text in heading_numbers if item_number == number]
        for number, count in Counter(number for number, _, _ in heading_numbers).items()
        if count > 1
    }
    caption_ids = []
    for location, text in captions:
        match = re.match(r"^([图表]\s*\d+[\-—]\d+)", text)
        if match:
            caption_ids.append((re.sub(r"\s+", "", match.group(1)).replace("—", "-"), location, text))
    duplicate_captions = {
        identifier: [text for item_id, _, text in caption_ids if item_id == identifier]
        for identifier, count in Counter(item_id for item_id, _, _ in caption_ids).items()
        if count > 1
    }
    wrong_chapter_captions = [
        {"location": location, "text": text}
        for location, text in captions
        if location.startswith("paragraph:") and "图5-2 业务流程图" in text
    ]

    structure_lines = [
        "# 原稿结构与数字主张审计",
        "",
        f"- 正文段落：{len(document.paragraphs)}",
        f"- 表格：{len(document.tables)}",
        f"- 内嵌图片：{len(document.inline_shapes)}",
        f"- 含数值主张：{len(rows)}",
        f"- D级主张：{sum(row['evidence_grade'] == 'D' for row in rows)}",
        "",
        "## 标题编号冲突",
        "",
    ]
    if duplicate_headings:
        for number, texts in duplicate_headings.items():
            structure_lines.append(f"- {number}：{'；'.join(texts)}")
    else:
        structure_lines.append("- 未发现重复标题编号。")
    structure_lines.extend(["", "## 图表编号冲突", ""])
    if duplicate_captions:
        for identifier, texts in duplicate_captions.items():
            structure_lines.append(f"- {identifier}：{'；'.join(texts)}")
    else:
        structure_lines.append("- 未发现重复题注编号。")
    for item in wrong_chapter_captions:
        structure_lines.append(f"- 第六章误用第五章图号：{item['text']}")
    (ARTIFACTS / "structure_audit.md").write_text("\n".join(structure_lines) + "\n", encoding="utf-8")

    print(
        json.dumps(
            {
                "claims": len(rows),
                "grade_d": sum(row["evidence_grade"] == "D" for row in rows),
                "duplicate_headings": duplicate_headings,
                "duplicate_captions": duplicate_captions,
                "wrong_chapter_captions": wrong_chapter_captions,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
