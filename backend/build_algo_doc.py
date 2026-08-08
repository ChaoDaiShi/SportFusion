#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate Algorithm Documentation DOCX for SportFusion."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import os

OUT = Path('../体融识界·SportFusion——算法技术文档.docx')

doc = Document()

# ── Style helpers ──────────────────────────────────────────────
def set_cn_font(run, name='宋体', size=12, bold=False):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def H1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(24)
    p.space_after = Pt(12)
    r = p.add_run(text); set_cn_font(r, '黑体', 16, True)
    p.style = doc.styles['Heading 1']

def H2(text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    r = p.add_run(text); set_cn_font(r, '黑体', 14, True)
    p.style = doc.styles['Heading 2']

def H3(text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    r = p.add_run(text); set_cn_font(r, '黑体', 12, True)
    p.style = doc.styles['Heading 3']

def P(text, sz=12, indent=False):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text); set_cn_font(r, '宋体', sz)

def Code(text):
    """Monospace code block"""
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = 'Consolas'

def Formula(text):
    """Centered formula"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    r = p.add_run(text); set_cn_font(r, 'Times New Roman', 11)

def TableRow(headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacer

# ================================================================
# TITLE PAGE
# ================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('体融识界·SportFusion'); set_cn_font(r, '黑体', 28, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('算法技术文档'); set_cn_font(r, '黑体', 24, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('基于NLP文本识别与多维度加权的\n多元经营企业体育业务边界识别与产业规模测算'); set_cn_font(r, '宋体', 14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('版本 2.0  —  2026年7月'); set_cn_font(r, '宋体', 11)

doc.add_page_break()

# ================================================================
# CHAPTER 1: OVERVIEW
# ================================================================
H1('第一章  算法体系概述')

H2('1.1 技术架构总览')
P('体融识界（SportFusion）的核心算法体系围绕"从企业级文本到业务线级识别，从二值判定到连续占比估算"两条主线构建，形成了一套三层递进、逐级深化的技术架构：', indent=True)
P('•  底层 — 文本预处理与特征工程层：中文分词（jieba）、体育自定义词典构建（271词/9大业态）、行业代码双层映射（36个代码）、多维特征提取（文本/体育/业务/代码四类特征）。')
P('•  中层 — 业务识别与分类层：业务线解析算法（Business Line Parser）将文本拆分为独立业务线，双通道融合识别模型（文本通道 + 行业代码通道）实现跨界经营精准判定。')
P('•  上层 — 比重测算与统计分析层：SportRatio 四维度加权模型估算体育业务连续占比值，置信度分层框架支撑差异化统计决策，跨界经营类型体系输出结构化产业洞察。')

H2('1.2 算法流水线')
P('完整算法流水线表示为：', indent=True)
Formula('原始工商数据 → 数据清洗 → jieba分词+体育词典匹配 → 业务线解析(BLP)')
Formula('→ 逐线体育分类 → 双通道融合识别 → SportRatio比重测算 → 置信度评估 → 跨界类型判定 → 产业统计输出')
P('单条企业文本的平均处理时间为 0.16 毫秒，全量 76,687 家企业仅需约 12 秒完成全部算法流水线。')

H2('1.3 核心算法一览')
TableRow(
    ['编号', '算法名称', '英文名称', '核心功能', '输入', '输出'],
    [
        ['A1', '业务线解析算法', 'Business Line Parser\n(BLP)', '将企业"主要业务活动"文本\n拆分为独立业务线', '自由文本', '业务线列表\n(≈3条/企业)'],
        ['A2', '双通道融合识别模型', 'Dual-Channel Fusion\nRecognition Model', '结合NLP语义+NACE代码\n判定体育业务归属', '业务线+行业代码', 'is_sport, confidence\ncrossover_type'],
        ['A3', '多维度比重测算模型', 'SportRatio Estimation\nModel', '在无营收数据条件下\n估算体育业务连续占比', '文本+代码特征', 'sport_ratio ∈ [0,1]'],
        ['A4', '跨界经营类型判定', 'Crossover Type\nClassifier', '基于代码与文本交叉分析\n识别三种跨界模式', '识别结果+行业代码', '纯跨界/潜在跨界\n/多元经营'],
    ]
)

doc.add_page_break()

# ================================================================
# CHAPTER 2: BUSINESS LINE PARSER (A1)
# ================================================================
H1('第二章  业务线解析算法（A1）')

H2('2.1 问题定义')
P('在多元经营背景下，一家企业的"主要业务活动"字段往往包含多种性质迥异的业务描述。例如，某企业登记的业务活动为：', indent=True)
P('   "体育赛事运营，广告设计制作，体育用品销售，企业管理咨询"', indent=True)
P('该企业同时经营体育赛事、广告、体育用品、咨询四类业务。传统方法将其整体归类为"体育企业"或"非体育企业"，无法精确区分其中的体育与非体育成分。', indent=True)

H2('2.2 算法设计')
P('业务线解析算法的设计灵感来源于自然语言处理中的句子分割（Sentence Segmentation）技术，借鉴了中文标点符号作为语义边界标记的语言学特性。算法将企业业务活动文本视为由多个独立业务描述构成的"微型文档"，通过识别业务分隔符将其分解为原子级别的业务线单元。', indent=True)

H3('2.2.1 步骤一：文本分割（Text Segmentation）')
P('使用正则表达式识别中文业务描述中常见的分隔符模式：', indent=True)
Code('SEPARATOR_PATTERN = r"[，,;；、/／\\n\\r。；；．\\.\\s]+"')
P('该模式覆盖了以下业务分隔符类别：', indent=True)
P('•  标点分隔符：中文逗号（，）、分号（；）、顿号（、）、句号（。）')
P('•  符号分隔符：斜杠（/）、反斜杠（\）、英文逗号（,）、分号（;）')
P('•  空白分隔符：空格、制表符、换行符')
P('分割后过滤长度小于 2 字符的无效片段（单独的符号、数字等），并对完全相同的业务线去重（保持首次出现顺序）。', indent=True)

P('实证统计：76,687 家企业的"主要业务活动"文本经分割，共解析出约 23 万条独立业务线，平均每家企业约 3 条，业务线数量分布在 1-15 条之间，中位数为 2 条。', indent=True)

H3('2.2.2 步骤二：逐线分类（Line-wise Classification）')
P('对每条业务线独立执行以下分类运算：', indent=True)
P('(a) jieba 分词：使用加载了 271 词体育自定义词典的 jieba 分词器对业务线文本进行分词。自定义词典的加载确保"电子竞技赛事"被正确切分为一个完整词条，而非"电子/竞技/赛事"三个碎片。', indent=True)
P('(b) 关键词匹配：将分词结果与体育词典进行精确匹配。词典使用 Set 数据结构实现 O(1) 查找，单条业务线的匹配耗时可忽略不计。', indent=True)
P('(c) 业态归类：通过反向索引（SPORT_WORD_TO_CATEGORY）查询命中关键词所属的业态类别，采用多数投票策略确定该业务线的主要业态归属。', indent=True)
P('(d) 匹配评分：score = min(len(keywords) / 3.0, 1.0)，即命中 3 个及以上体育关键词即为满分 1.0，命中 1 个为 0.33，命中 2 个为 0.67。', indent=True)

H3('2.2.3 步骤三：边界聚合（Boundary Aggregation）')
P('汇总所有业务线的分类结果，形成企业的"体育业务边界全景图"——精确列出：', indent=True)
P('•  哪些业务线属于体育业务（及各自的业态类型）')
P('•  哪些业务线不属于体育业务')
P('•  体育业务线占全部业务线的比例（即 W1 业务范围占比）')
P('边界聚合的输出被直接输入到下游的双通道融合识别模型和 SportRatio 测算模型中，构成了整个算法体系的粒度基础。', indent=True)

H2('2.3 算法复杂度')
TableRow(
    ['指标', '数值', '说明'],
    [
        ['时间复杂度', 'O(n·m)', 'n=企业数, m=平均业务线数(≈3)'],
        ['单条处理时间', '0.16 ms', 'jieba分词(0.13ms) + 词典匹配(0.03ms)'],
        ['全量处理时间', '≈12 秒', '76,687家企业'],
        ['吞吐率', '6,325 条/秒', '远超年度普查的实时性需求'],
        ['空间复杂度', 'O(d)', 'd=词典词条数(271)'],
    ]
)

doc.add_page_break()

# ================================================================
# CHAPTER 3: DUAL-CHANNEL FUSION MODEL (A2)
# ================================================================
H1('第三章  双通道融合识别模型（A2）')

H2('3.1 设计思想')
P('双通道融合识别模型的核心设计思想是"语义理解 + 结构约束"的互补融合：', indent=True)
P('•  文本分析通道（语义层面）：基于NLP技术从业务活动文本中提取体育语义信号，灵活捕捉各类体育业务的文本特征，不受行业代码分类体系的固有框架限制。')
P('•  行业代码通道（结构层面）：利用GB/T 4754-2017行业分类体系的结构性约束，为识别结果提供与现行统计制度相兼容的"锚点"，确保识别结果可被现行统计体系直接采纳。')
P('两个通道通过四层置信度规则进行加权融合，实现优势互补——文本通道负责"发现"，代码通道负责"锚定"，融合结果兼顾灵敏性与可靠性。', indent=True)

H2('3.2 行业代码双层映射体系')
P('在融合之前，首先建立了 GB/T 4754-2017 与体育产业的双层代码映射关系，为文本分析结果与行业分类体系之间搭建标准化的桥梁。', indent=True)

TableRow(
    ['层级', '代码类型', '数量', '判定逻辑', '示例代码'],
    [
        ['第一层', '直接体育代码', '19个', '企业主营业务即为体育，\n行业代码本身即强烈信号', '8911(体育组织)\n8930(健身休闲)\n2441(球类制造)'],
        ['第二层', '间接相关代码', '17个', '部分企业可能涉及体育，\n需文本辅助确认', '8399(未列明教育)\n7259(广告服务)\n6513(软件开发)'],
    ]
)

P('共计 36 个体育相关代码，覆盖了体育赛事、健身休闲、体育用品、体育培训、体育场馆、体育传媒、体育管理、体育科技、体育旅游等业态。', indent=True)

H2('3.3 四层置信度规则')
P('模型将文本分析结果与行业代码类型进行交叉组合，定义了四层置信度判定规则，形成从"几乎确定"到"纯文本推断"的梯度置信体系：', indent=True)

TableRow(
    ['规则', '行业代码', '文本匹配', '置信度', '识别场景', '企业数量\n(76,687全量)'],
    [
        ['规则1\n(最高)', '直接体育代码\n(19个)', '✓ 命中', '0.95', '双重确认：代码+文本\n均指向体育', '≈5,286家\n(58.6%)'],
        ['规则2\n(高)', '直接体育代码\n(19个)', '✗ 未命中', '0.75', '代码明确，文本简略\n或使用了词典外表述', '≈2,730家\n(30.3%)'],
        ['规则3\n(中)', '间接相关代码\n(17个)', '✓ 命中', '0.70', '跨界经营确认：\n文本证实代码未明示\n的体育业务', '≈56家\n(0.6%)'],
        ['规则4\n(中)', '无关代码\n(其余所有)', '✓ 命中', '0.60', '纯文本发现：完全依赖\nNLP能力识别跨界\n（传统方法无此能力）', '≈951家\n(10.5%)'],
    ]
)

H2('3.4 判定阈值')
P('体育业务占比（sport_ratio）≥ 0.10 的企业被判定为体育企业。该阈值经与直接体育代码企业的交叉验证确定：', indent=True)
P('•  阈值过低（< 0.05）：大量仅有一条模糊体育相关描述的企业被误判为体育企业，精确率下降。')
P('•  阈值过高（> 0.20）：部分体育业务占比不高但确实从事体育经营的企业被漏检，召回率下降。')
P('•  0.10 阈值在精确率（Precision）与召回率（Recall）之间取得了经验最优平衡，且与实际统计工作中"主营业务占比超过10%计为产业组成部分"的惯例一致。', indent=True)

H2('3.5 模型输出')
P('对于每家企业，模型输出包含以下结构化信息的识别结果对象：', indent=True)
TableRow(
    ['输出字段', '类型', '说明'],
    [
        ['is_sport', 'bool', '是否判定为体育企业'],
        ['sport_category', 'str', '主要体育业态（9分类之一）'],
        ['sport_ratio', 'float', '体育业务占比 [0, 1]'],
        ['confidence', 'float', '判定置信度 [0, 1]'],
        ['is_crossover', 'bool', '是否属于跨界经营'],
        ['crossover_type', 'str', '跨界类型标签'],
        ['code_type', 'str', '行业代码类型 (direct/indirect/none)'],
        ['keywords', 'list[str]', '命中的体育关键词证据链'],
        ['feature_weights', 'dict', 'W1-W4各维度得分'],
    ]
)

doc.add_page_break()

# ================================================================
# CHAPTER 4: SPORTRATIO MODEL (A3)
# ================================================================
H1('第四章  SportRatio 多维度比重测算模型（A3）')

H2('4.1 问题定义')
P('产业规模测算的核心难题在于：企业工商登记数据中不含营收信息，无法直接获取体育业务的财务占比。传统的行业代码法仅输出二值结果（是/否体育企业），将所有体育企业等权计入产业规模，无法反映体育业务在企业经营中的实际比重差异。', indent=True)
P('SportRatio 模型的设计目标：在零财务数据的约束条件下，利用文本特征和代码先验信息，对每家企业输出一个连续的体育业务占比估值（sport_ratio ∈ [0, 1]），为产业规模的比例测算提供方法基础。', indent=True)

H2('4.2 模型架构')
P('SportRatio 是一个四维度线性加权模型，综合四个可独立计算的代理指标来逼近真实的体育业务占比：', indent=True)
Formula('SportRatio = 0.40·W1 + 0.25·W2 + 0.25·W3 + 0.10·W4')
Formula('(截断至 [0, 1] 区间)')

H3('4.2.1 W1 — 业务范围占比（权重 0.40）')
P('定义：W1 = sport_business_lines / total_business_lines', indent=True)
P('含义：直接反映体育业务线在企业登记的全部业务线中的占比。例如，某企业有 5 条业务线，其中 2 条属于体育，则 W1 = 0.40。', indent=True)
P('合理性基础：企业在工商登记中填报的业务活动通常按重要性递减排序，体育业务线数量的占比与企业实际的体育投入分配存在统计学上的正相关关系。该维度权重最高（0.40），因为它最直接地度量了体育在企业经营范围中的"空间占比"。', indent=True)

H3('4.2.2 W2 — 关键词密度（权重 0.25）')
P('定义：W2 = min(sport_keyword_count / total_tokens × 10, 1.0)', indent=True)
P('含义：将体育关键词命中数除以总词数再乘以归一化因子 10，反映了体育业务在文本描述中的"语义密度"。描述越详细、越聚焦于体育的企业，W2 越高。', indent=True)
P('示例："体育赛事运营，专业马拉松组织，运动品牌推广"（3个体育关键词，总词数6） → W2 = 3/6×10 = 5.0 → 截断为 1.0；"体育用品销售"（1个体育关键词，3个总词） → W2 = 1/3×10 ≈ 0.33。', indent=True)

H3('4.2.3 W3 — 行业代码权重（权重 0.25）')
P('定义：W3 ∈ {0.85, 0.30, 0.00}，取决于行业代码类型：', indent=True)
TableRow(
    ['代码类型', 'W3 值', '先验含义'],
    [
        ['直接体育代码', '0.85', '以体育为主业的高先验概率，即代码本身就强烈暗示企业主要从事体育业务'],
        ['间接相关代码', '0.30', '可能含体育的中等先验概率，需文本进一步确认'],
        ['无关代码', '0.00', '无先验信息，完全依赖文本分析（纯跨界场景）'],
    ]
)
P('该维度将行业分类体系的结构性信息以先验概率的形式引入模型，在缺乏营收真值的条件下提供了重要的"锚定"信息。', indent=True)

H3('4.2.4 W4 — 业态覆盖度（权重 0.10）')
P('定义：W4 = sport_category_count / 9（9 为总业态数）', indent=True)
P('含义：衡量企业体育业务的多元化程度。同时涉及多个体育业态（如赛事+培训+用品）的企业，其体育投入通常比单一业态的企业更深入、更系统，给予更高的综合权重。', indent=True)
P('该维度权重最低（0.10），仅作为微调因子，因为业态多样性与业务深度之间的正相关并非必然——少数企业可能"浅尝辄止"地涉及多个业态。', indent=True)

H2('4.3 权重设计原理')
P('权重分配遵循"文本内容为主（0.65 = W1+W2），代码先验为辅（0.25 = W3），多元程度微调（0.10 = W4）"的设计原则。这一分配策略基于以下考量：', indent=True)
P('(1) 文本信息是最直接反映企业实际业务的信号来源，赋予最高综合权重（65%）。', indent=True)
P('(2) 行业代码虽对体育主业企业具有强识别力，但在跨界场景下识别力为零，作为辅助信息赋予 25% 权重。', indent=True)
P('(3) 业态多样性的正相关性较弱且易受噪声影响，仅作为微调因子赋予 10% 权重。', indent=True)

H2('4.4 模型局限与改进方向')
P('当前 SportRatio 属于无监督线性加权模型，权重基于经验设定，虽经交叉验证但缺乏营收真值的直接校准。未来若可获取部分企业的营收明细数据（如通过税务系统对接），可通过以下路径升级：', indent=True)
P('(1) 有监督回归：以营收分成为标签，使用梯度提升树（XGBoost）或神经网络学习 W1-W4 的最优组合权重及非线性交互项。', indent=True)
P('(2) 多源数据增强：融入社保缴纳数据（人员投入规模）、知识产权数据（技术创新方向）、网络信息（业务动态更新）等，构建更高维度的特征空间。', indent=True)

doc.add_page_break()

# ================================================================
# CHAPTER 5: CROSSOVER TYPE CLASSIFICATION (A4)
# ================================================================
H1('第五章  跨界经营类型识别（A4）')

H2('5.1 问题定义')
P('跨界经营是多元经营背景下产业统计的核心挑战。传统统计方法对跨界企业完全"失明"——行业代码法仅能依据注册代码判定，无法发现代码之外的体育业务。SportFusion 通过交叉分析行业代码类型与文本识别结果，建立了三种跨界经营类型的系统分类体系。', indent=True)

H2('5.2 三种跨界类型')
TableRow(
    ['类型', '判定条件', '置信度', '识别数量\n(76,687)', '典型场景'],
    [
        ['纯跨界\n(Pure\nCrossover)', '行业代码与体育无关\n+ 文本命中体育关键词', '0.60', '951家\n(增量100%)', '软件开发公司\n兼营电子竞技赛事运营'],
        ['潜在跨界\n(Potential\nCrossover)', '行业代码间接相关\n+ 文本确认含体育业务', '0.70', '56家', '广告公司兼营\n体育赛事推广'],
        ['多元经营\n(Diversified\nOperation)', '行业代码为直接体育\n+ 存在多条非体育业务线', '0.85', '≈7,943家\n(含非体育线)', '体育用品制造企业\n兼营物流/贸易'],
    ]
)

H2('5.3 判定逻辑')
P('跨界类型判定采用以下决策树逻辑（伪代码）：', indent=True)
Code('if is_sport and code_type == "none":')
Code('    crossover_type = "纯跨界（行业代码非体育，文本有体育业务）"')
Code('elif is_sport and code_type == "indirect":')
Code('    crossover_type = "潜在跨界（间接行业代码，文本有体育业务）"')
Code('elif is_sport and code_type == "direct" and non_sport_count > 0:')
Code('    crossover_type = f"多元经营（体育+{non_sport_count}条非体育业务）"')
Code('else:')
Code('    crossover_type = ""  # 纯体育企业，非跨界')

H2('5.4 关键发现')
P('经 76,687 家企业全量实证，跨界经营类型分析揭示了以下关键统计规律：', indent=True)
P('•  纯跨界企业（951家）全部由 NLP 文本通道独立发现，传统行业代码法对此类企业的识别率为 0%。这 951 家企业构成了传统统计视野之外的"隐形产业带"。', indent=True)
P('•  新兴业态的跨界率极高：电子竞技 91%、体育传媒 94%，意味着这些业态几乎完全依赖文本分析才得以发现。成熟业态的跨界率极低：健身休闲 2%、体育用品 7%，被现行行业分类体系较好覆盖。', indent=True)
P('•  "纯跨界"与"潜在跨界"合计 1,007 家企业，占全部识别体育企业的 11.3%，纠正了约 6 亿元规模的产业体量低估。', indent=True)

doc.add_page_break()

# ================================================================
# CHAPTER 6: KEY TECHNOLOGY DECISIONS
# ================================================================
H1('第六章  关键技术选择分析')

H2('6.1 规则匹配 vs 深度学习')
P('本研究在文本识别环节采用了"规则匹配（Rule-based Matching）+ TF-IDF 关键词提取"的组合方案，而未直接采用深度学习模型（如 BERT 微调）。这一选择基于以下严谨技术考量，也是本项目工程方法论的核心决策点：', indent=True)

TableRow(
    ['考量维度', '规则匹配方案\n(jieba + 自定义词典)', '深度学习方案\n(BERT 微调)', '本项目选择'],
    [
        ['标注数据需求', '零标注数据\n仅需构建和维护\n271词体育词典', '需1,000-5,000条\n人工标注训练数据\n（成本高、标准难统一）', '规则匹配 ✓\n（冷启动可行）'],
        ['计算效率', '6,325条/秒\n全量12秒完成', '约50-500条/秒(GPU)\n全量需数十分钟\n至数小时', '规则匹配 ✓\n（实时响应）'],
        ['可解释性', '每条结果可追溯\n到具体关键词证据链', '注意力权重可部分\n解释但不够直观\n业务人员难以理解', '规则匹配 ✓\n（统计透明度）'],
        ['维护更新', '"热更新"：添加\n词条即可，无需重训练', '需重新收集标注数据\n重新训练模型\n（流程长、成本高）', '规则匹配 ✓\n（运维灵活性）'],
        ['语义理解深度', '仅能匹配显式关键词\n无法理解隐含语义\n（如"组织群众体育活动"）', '可理解上下文语境\n捕捉隐性体育信号\n泛化能力强', '深度学习 →\n（未来升级路径）'],
    ]
)

H2('6.2 未来升级路径')
P('本研究充分认识到规则匹配方法在语义理解深度方面的局限性，并在技术架构中预留了深度学习升级接口。建议按以下两阶段路径推进：', indent=True)
P('第一阶段（短期）：维持当前规则匹配方案作为主引擎，利用其已积累的识别结果进行"伪标签"采样，初步构建标注数据集（目标：1,000-2,000 条高质量标注）。', indent=True)
P('第二阶段（中期）：基于累积的标注数据，采用领域自适应预训练策略（Domain-Adaptive Pretraining），在 BERT-base-Chinese 基础上注入体育产业语料进行持续预训练，然后微调构建体育文本分类器。最终目标是与规则匹配形成"双引擎"架构——规则引擎处理高置信度场景，深度学习引擎处理语义复杂/规则漏检场景。', indent=True)

doc.add_page_break()

# ================================================================
# CHAPTER 7: PERFORMANCE SUMMARY
# ================================================================
H1('第七章  算法性能总览')

H2('7.1 关键性能指标')
TableRow(
    ['指标', '数值', '对比基准', '提升幅度'],
    [
        ['识别覆盖率', '11.67%\n(8,950家/76,687家)', '传统行业代码法\n10.45% (8,016家)', '+11.7%'],
        ['跨界发现企业数', '1,007家\n(纯跨界951+潜在跨界56)', '传统行业代码法\n0家（无法识别）', '从0到1,007\n（突破性）'],
        ['新兴业态覆盖率', '电子竞技 43家(91%跨界)\n体育传媒 16家(94%跨界)', '传统法 <5%\n几乎完全遗漏', '>90%'],
        ['人工核查缩减', '81%\n(5,000→951家)', '全量人工核查\n5,000家', '效率提升\n5.3×'],
        ['年统计成本节约', '≈40万元', '纯人工统计流程', '成本降低\n64.8%'],
        ['单条处理延迟', '0.16 ms', 'BERT推理\n20-200 ms(GPU)', '125-1,250×\n速度优势'],
        ['全量处理时间', '≈12 秒\n(76,687家企业)', 'BERT(GPU)\n数十分钟-数小时', '50-100×\n速度优势'],
        ['产业体量纠正', '≈6亿元', '传统法遗漏\n934家企业', '统计学显著'],
    ]
)

H2('7.2 算法复杂度汇总')
TableRow(
    ['算法模块', '时间复杂度', '空间复杂度', '可扩展性'],
    [
        ['中文分词+词典匹配', 'O(n)', 'O(d), d=271', '线性扩展\n可处理百万级'],
        ['业务线解析(BLP)', 'O(n·m), m≈3', 'O(1)', '近似线性'],
        ['双通道融合识别', 'O(n)', 'O(1)', '线性扩展'],
        ['SportRatio比重测算', 'O(n·m)', 'O(1)', '近似线性'],
        ['全流水线(端到端)', 'O(n·m)', 'O(d)', '6,325条/秒吞吐率'],
    ]
)

H2('7.3 代码文件索引')
TableRow(
    ['文件路径', '功能', '核心类/函数'],
    [
        ['backend/utils/text_tokenizer.py', '文本分词+体育词典', 'tokenize(), match_sport_keywords()\nmatch_sport_by_category()'],
        ['backend/utils/industry_code.py', '行业代码映射体系', 'get_code_type(), is_direct_sport_code()\nclassify_by_text_and_code()'],
        ['backend/services/sport_recognition.py', '核心识别算法', 'parse_business_lines()\nclassify_business_line()\ncalculate_sport_ratio()\nrecognize_sport_business()'],
        ['backend/services/nlp_preprocess.py', 'NLP预处理+特征工程', 'preprocess_enterprise()\nbatch_preprocess_enterprises()'],
    ]
)

# ── Save ───────────────────────────────────────────────────────
doc.save(str(OUT))
print(f'Algorithm documentation saved to: {OUT}')
print('Done.')
