"""Verify the generated SportFusion visual report and its evidence artifacts."""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageChops
from pypdf import PdfReader

EXPECTED_SOURCE_SHA256 = "8b403f4bf3f6325761acbf028dc319f76d7a20df10323b87625da4bca7586682"
EXPECTED_SOURCE_SIZE = 22_886_541
EXPECTED_COUNTS = {
    "paragraphs": 417,
    "tables": 56,
    "inline_shapes": 28,
    "sections": 2,
}
CORE_VALUES = ["76,687", "8,016", "8,950", "934", "977", "6,220", "2,730", "2,170.80"]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def compact(value: str) -> str:
    return "".join(value.split())


def load_builder(path: Path):
    spec = importlib.util.spec_from_file_location("sportfusion_visual_report_builder", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot import builder: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--manifest", required=True, type=Path)
    parser.add_argument("--pdf", required=True, type=Path)
    parser.add_argument("--builder", required=True, type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    builder = load_builder(args.builder)
    manifest = json.loads(args.manifest.read_text(encoding="utf-8"))

    assert args.source.stat().st_size == EXPECTED_SOURCE_SIZE
    assert sha256(args.source) == EXPECTED_SOURCE_SHA256
    assert args.output.exists()
    assert sha256(args.output) == manifest["report"]["output_sha256"]
    assert manifest["report"]["source_sha256"] == EXPECTED_SOURCE_SHA256

    with zipfile.ZipFile(args.output) as archive:
        assert archive.testzip() is None

    document = Document(args.output)
    counts = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "sections": len(document.sections),
    }
    assert counts == EXPECTED_COUNTS
    assert manifest["report"]["final_counts"] == EXPECTED_COUNTS

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    full_text = "\n".join(paragraphs + table_text)
    for value in CORE_VALUES:
        assert value in full_text, value
    assert "Backend：269" not in full_text
    assert "Frontend：32" not in full_text
    assert "Backend：聚焦48/48；全量628 pass、3 skip、1 xfail、2 fail；Frontend：34/34+build" in full_text
    assert "待补充图4-2" not in full_text

    expected_captions = {
        f"图4-2（{item['letter']}）{item['title']}" for item in builder.SCREENSHOTS
    }
    assert len(expected_captions) == 9
    assert sum(paragraph in expected_captions for paragraph in paragraphs) == 9

    screenshots = manifest["screenshots"]
    assert len(screenshots) == 9
    for screenshot in screenshots:
        raw_path = Path(screenshot["raw"])
        processed_path = Path(screenshot["processed"])
        assert raw_path.exists() and processed_path.exists()
        assert sha256(raw_path) == screenshot["raw_sha256"]
        assert sha256(processed_path) == screenshot["processed_sha256"]
        assert screenshot["raw_size"] == [1920, 1080]
        assert screenshot["raw_pixels_preserved"] is True
        raw = Image.open(raw_path).convert("RGB")
        processed = Image.open(processed_path).convert("RGB")
        left, top, right, bottom = screenshot["raw_pixel_rectangle"]
        assert ImageChops.difference(raw, processed.crop((left, top, right, bottom))).getbbox() is None

    reader = PdfReader(args.pdf)
    assert len(reader.pages) == 70
    for page in reader.pages:
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        assert abs((width / height) - (210 / 297)) < 0.01

    pdf_pages = [compact(page.extract_text() or "") for page in reader.pages]
    for toc_entry in builder.TOC_PAGE_UPDATES.values():
        heading, printed_page_text = toc_entry.rsplit("\t", 1)
        pdf_page_index = int(printed_page_text) + 2
        assert compact(heading) in pdf_pages[pdf_page_index], (heading, printed_page_text)

    summary = {
        "source_sha256": sha256(args.source),
        "output_sha256": sha256(args.output),
        "output_bytes": args.output.stat().st_size,
        "docx_counts": counts,
        "screenshots": len(screenshots),
        "raw_pixels_preserved": all(item["raw_pixels_preserved"] for item in screenshots),
        "pdf_pages": len(reader.pages),
        "toc_entries_verified": len(builder.TOC_PAGE_UPDATES),
    }
    print(json.dumps(summary, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
