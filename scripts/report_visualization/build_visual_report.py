"""Build the audited SportFusion system-visualization report.

The script preserves the source document and creates a new DOCX. Raw screenshots are
kept pixel-identical inside a restrained frame; all labels sit outside the captured
system rectangle so the report never disguises demo or empty states as real results.
"""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageChops, ImageDraw, ImageFont

SCREENSHOTS = [
    {
        "id": "a-monitoring",
        "letter": "a",
        "title": "统计监测驾驶舱：五步工作流与风险总览",
        "badge": "演示界面",
        "badge_color": "#3158D8",
        "note": "界面状态：功能演示。页面已显示“演示数据保障”；卡片数值仅用于界面联调，不作为第三章实证数据来源。",
    },
    {
        "id": "b-data",
        "letter": "b",
        "title": "企业数据治理：文件导入与过程追踪入口",
        "badge": "真实空态",
        "badge_color": "#168B78",
        "note": "界面状态：真实空态。截图展示未上传文件时的系统初始状态，不含模拟结果。",
    },
    {
        "id": "c-recognition",
        "letter": "c",
        "title": "企业边界识别：文本—代码双通道识别入口",
        "badge": "真实空态",
        "badge_color": "#168B78",
        "note": "界面状态：真实空态。截图展示单家与批量识别入口，未填入企业信息，未生成识别结论。",
    },
    {
        "id": "d-share",
        "letter": "d",
        "title": "经营比重测算：SportShare结果与人工校准入口",
        "badge": "真实空态",
        "badge_color": "#168B78",
        "note": "界面状态：真实空态。当前运行未加载正式批次与SportShare结果，页面中的0为未测算状态，不代表实证样本为0。",
    },
    {
        "id": "e-scale",
        "letter": "e",
        "title": "产业规模测算：宏观总量校准与旧路径隔离",
        "badge": "正式口径",
        "badge_color": "#277A3E",
        "note": "界面状态：正式口径说明页。2,170.80亿元来自项目锁定的官方总量配置；页面尚未执行本次规模测算，不新增实证结果。",
    },
    {
        "id": "f-evaluation",
        "letter": "f",
        "title": "模型性能评估：识别效果、异常输入与资源消耗",
        "badge": "演示界面",
        "badge_color": "#3158D8",
        "note": "界面状态：功能演示。页面已明确提示“当前显示演示评测数据”；其中指标不作为本文验证集或Golden结果。",
    },
    {
        "id": "g-review",
        "letter": "g",
        "title": "人工复核工作台：优先级、双人复核与仲裁状态",
        "badge": "功能演示",
        "badge_color": "#A15D08",
        "note": "界面状态：功能演示。4条“演示样本”由真实复核接口生成，用于展示待分配、复核中、已确认和待仲裁状态，不进入正式名录。",
    },
    {
        "id": "h-directory",
        "letter": "h",
        "title": "动态企业名录：finalized准入与筛选入口",
        "badge": "真实空态",
        "badge_color": "#168B78",
        "note": "界面状态：真实空态。当前运行没有已确认或锁定的正式企业，因此名录为空；待复核与争议任务不会被静默纳入。",
    },
    {
        "id": "i-export",
        "letter": "i",
        "title": "报告与成果中心：版本锁定与结构化导出",
        "badge": "演示界面",
        "badge_color": "#3158D8",
        "note": "界面状态：功能演示。页面已显示“演示数据保障”；导出卡片用于说明成果类型，不证明相应文件已按正式批次生成。",
    },
]

TOC_PAGE_UPDATES = {
    "4.2年度体育产业名录更新\t39": "4.2年度体育产业名录更新\t45",
    "4.3跨界体育经营主体专项核查\t40": "4.3跨界体育经营主体专项核查\t46",
    "4.4区域产业结构与政策研判\t41": "4.4区域产业结构与政策研判\t47",
    "4.5部门协同与常态化统计监测\t42": "4.5部门协同与常态化统计监测\t48",
    "4.6应用价值与经济社会效益\t43": "4.6应用价值与经济社会效益\t49",
    "4.7推广应用路径\t44": "4.7推广应用路径\t50",
    "第五章结论、创新贡献与政策建议\t45": "第五章结论、创新贡献与政策建议\t51",
    "5.1主要成果\t45": "5.1主要成果\t51",
    "5.2方法创新与实践贡献\t45": "5.2方法创新与实践贡献\t51",
    "5.3与传统方法的综合比较\t45": "5.3与传统方法的综合比较\t52",
    "5.4研究局限与改进方向\t46": "5.4研究局限与改进方向\t52",
    "5.5统计监测优化与政策建议\t46": "5.5统计监测优化与政策建议\t53",
    "参考文献\t48": "参考文献\t55",
    "附录A数据口径与字段说明\t50": "附录A数据口径与字段说明\t57",
    "附录B验证与异常输入明细\t52": "附录B验证与异常输入明细\t59",
    "附录C全样本与规模结果\t54": "附录C全样本与规模结果\t61",
    "附录D复核任务与名录字段\t56": "附录D复核任务与名录字段\t63",
    "附录E系统运行与复现说明\t58": "附录E系统运行与复现说明\t65",
    "附录F电子支撑材料清单\t59": "附录F电子支撑材料清单\t67",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    names = ["msyhbd.ttc" if bold else "msyh.ttc", "simhei.ttf"]
    for name in names:
        candidate = Path("C:/Windows/Fonts") / name
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    left, _, right, _ = draw.textbbox((0, 0), text, font=font)
    return right - left


def create_framed_screenshot(raw_path: Path, output_path: Path, item: dict) -> dict:
    raw = Image.open(raw_path).convert("RGB")
    border = 8
    footer = 96
    canvas = Image.new("RGB", (raw.width + border * 2, raw.height + border * 2 + footer), "#F6F1E7")
    canvas.paste(raw, (border, border))

    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, canvas.width - 1, canvas.height - 1), outline="#D8CDB9", width=2)
    footer_y = border + raw.height
    draw.rectangle((border, footer_y, border + raw.width, footer_y + footer), fill="#FCFAF5")
    draw.rectangle((border, footer_y, border + 10, footer_y + footer), fill="#3158D8")

    title_font = load_font(31, bold=True)
    badge_font = load_font(24, bold=True)
    title = f"图4-2（{item['letter']}） {item['title']}"
    draw.text((border + 36, footer_y + 27), title, font=title_font, fill="#10212B")

    badge = item["badge"]
    badge_width = rounded_text_width(draw, badge, badge_font) + 46
    x2 = border + raw.width - 26
    x1 = x2 - badge_width
    y1, y2 = footer_y + 20, footer_y + 76
    draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=item["badge_color"])
    draw.text((x1 + 23, y1 + 12), badge, font=badge_font, fill="white")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output_path, format="PNG", optimize=True)

    check = Image.open(output_path).convert("RGB").crop((border, border, border + raw.width, border + raw.height))
    pixel_identity = ImageChops.difference(raw, check).getbbox() is None
    if not pixel_identity:
        raise RuntimeError(f"Raw screenshot pixels changed while framing: {raw_path}")

    return {
        "id": item["id"],
        "raw": str(raw_path),
        "processed": str(output_path),
        "raw_sha256": sha256(raw_path),
        "processed_sha256": sha256(output_path),
        "raw_size": list(raw.size),
        "processed_size": list(canvas.size),
        "raw_pixel_rectangle": [border, border, border + raw.width, border + raw.height],
        "raw_pixels_preserved": pixel_identity,
        "badge": item["badge"],
    }


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def ensure_note_style(document: Document) -> str:
    for candidate in ("图表注释", "图表题注", "Normal"):
        if candidate in document.styles:
            return candidate
    style = document.styles.add_style("图表注释", WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(9)
    return style.name


def add_page_break_before(anchor) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_screenshot_block(anchor, image_path: Path, item: dict, note_style: str) -> None:
    image_paragraph = anchor.insert_paragraph_before(style="正文无缩进" if "正文无缩进" in anchor.part.document.styles else None)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(3)
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.add_run().add_picture(str(image_path), width=Cm(15.4))

    caption = anchor.insert_paragraph_before(style="图表题注" if "图表题注" in anchor.part.document.styles else None)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run(f"图4-2（{item['letter']}）{item['title']}")
    set_east_asia_font(run, "宋体")

    note = anchor.insert_paragraph_before(style=note_style)
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(5)
    note.paragraph_format.keep_with_next = False
    run = note.add_run(item["note"])
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(89, 89, 89)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_analysis_paragraph(anchor, text: str, bold_prefix: str | None = None) -> None:
    paragraph = anchor.insert_paragraph_before(style="Normal")
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        set_east_asia_font(first, "宋体")
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_east_asia_font(rest, "宋体")
    else:
        run = paragraph.add_run(text)
        set_east_asia_font(run, "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_east_asia_font(run, "宋体")


def update_quality_gate_table(document: Document) -> int:
    """Synchronize the compact quality-gate table with the audited test run."""
    matches = 0
    for table in document.tables:
        for row in table.rows:
            if row.cells[0].text.strip() != "发布质量门":
                continue
            if len(row.cells) != 3:
                raise RuntimeError("Unexpected quality-gate table structure")
            replace_paragraph_text(
                row.cells[1].paragraphs[0],
                "Backend：聚焦48/48；全量628 pass、3 skip、1 xfail、2 fail；Frontend：34/34+build",
            )
            replace_paragraph_text(
                row.cells[2].paragraphs[0],
                "Ruff与锁检查通过；2项Alembic基线漂移待修复",
            )
            matches += 1
    if matches != 1:
        raise RuntimeError(f"Expected one quality-gate table row, found {matches}")
    return matches


def update_report(source: Path, output: Path, processed_dir: Path) -> dict:
    document = Document(source)
    original_counts = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "sections": len(document.sections),
    }

    toc_updates_applied = 0
    for paragraph in document.paragraphs:
        replacement = TOC_PAGE_UPDATES.get(paragraph.text)
        if replacement:
            new_page = replacement.rsplit("\t", 1)[1]
            if not paragraph.runs or not paragraph.runs[-1].text.startswith("\t"):
                raise RuntimeError(f"Manual TOC run structure changed: {paragraph.text}")
            paragraph.runs[-1].text = f"\t{new_page}"
            toc_updates_applied += 1
    if toc_updates_applied != len(TOC_PAGE_UPDATES):
        raise RuntimeError(
            f"Expected {len(TOC_PAGE_UPDATES)} manual TOC updates, applied {toc_updates_applied}"
        )

    caption_candidates = [p for p in document.paragraphs if p.text.startswith("图4-2 SportFusion Web系统演示路径")]
    if len(caption_candidates) != 1:
        raise RuntimeError(f"Expected one Figure 4-2 placeholder caption, found {len(caption_candidates)}")
    old_caption = caption_candidates[0]
    old_note_element = old_caption._p.getnext()
    old_note = next((p for p in document.paragraphs if p._p is old_note_element), None)
    if old_note is None or "最终提交前" not in old_note.text:
        raise RuntimeError("Figure 4-2 placeholder note was not found next to the caption")

    intro_candidates = [p for p in document.paragraphs if p.text.startswith("系统演示建议按照")]
    if len(intro_candidates) != 1:
        raise RuntimeError("Could not uniquely locate the system-demonstration introduction")
    replace_paragraph_text(
        intro_candidates[0],
        "系统演示按照“监测驾驶舱→数据管理→企业识别→经营比重→规模测算→模型评估→人工复核→动态名录→导出与锁定”的顺序展开。演示时重点说明SportScore与SportShare的职责分离、官方总量只在规模聚合阶段进入、SportShare来源标签与artifact_required错误态、P1—P4优先级、双人复核与仲裁、finalized名录过滤、批次锁定以及Provenance版本追溯。图4-2以同一提交版本的九张系统页面呈现上述链路；其中演示指标和演示任务均单独标注，不作为第三章实证数据来源。",
    )

    quality_candidates = [p for p in document.paragraphs if p.text.startswith("全链路质量与运行效率以当前发布候选")]
    if len(quality_candidates) != 1:
        raise RuntimeError("Could not uniquely locate the quality-gate paragraph")
    replace_paragraph_text(
        quality_candidates[0],
        "全链路质量与运行效率按当前发布候选的自动化质量门和严格性能基准分别记录。本次系统口径同步后，后端复核接口及Phase 4工作流聚焦测试48项全部通过；后端全量回归为628项通过、3项跳过、1项预期失败、2项Alembic基线漂移失败，后两项在主分支同样复现，尚未修复，不能表述为全量通过。前端34项测试和生产构建通过。对76,687条企业记录执行3次预热后连续完整运行5次的既有严格性能基准保持不变：中位耗时11.29秒，折合约0.147毫秒/条、约6,792条/秒。早期BATCH-20260803-R1记录的9.6秒单次运行仅作为历史参考，不作为当前跨设备稳定性能。",
    )
    quality_gate_rows_updated = update_quality_gate_table(document)

    note_style = ensure_note_style(document)
    add_page_break_before(old_caption)
    for index, item in enumerate(SCREENSHOTS):
        add_screenshot_block(old_caption, processed_dir / f"{item['id']}.png", item, note_style)
        if index % 2 == 1 and index != len(SCREENSHOTS) - 1:
            add_page_break_before(old_caption)

    add_analysis_paragraph(
        old_caption,
        "图4-2（a）—（i）把数据导入、双通道识别、SportShare估计、宏观校准规模分配、模型评估、人工复核、finalized名录准入和成果导出组织为一条可审计链路。页面分工与第二、三章方法口径一致：SportScore用于判断体育相关性，SportShare用于描述经营活动结构，官方总量只在聚合阶段约束九类业态和区域结构，复核结果则通过状态机控制是否进入名录。",
    )
    add_analysis_paragraph(
        old_caption,
        "可视化审查结论：九张截图均来自提交f2f17f8的同一运行版本，视口为1920×1080。监测驾驶舱、模型评估和成果中心属于系统自带演示界面；人工复核页使用4条明确命名的演示样本走通真实接口；数据治理、企业识别、SportShare和动态名录展示未载入正式批次时的真实空态；规模页展示正式宏观校准口径但未执行本次测算。上述页面证明功能结构和失败边界，不替代第三章的样本、标签、模型或规模证据。",
        bold_prefix="可视化审查结论：",
    )
    add_analysis_paragraph(
        old_caption,
        "内容一致性审查未改变第三章的实证口径：全样本仍为76,687家，传统直接体育代码覆盖8,016家，SportFusion候选8,950家，传统边界外补充934家，跨界或多元经营候选977家；SportShare来源仍为6,220家模型估计和2,730家分层回退，区域总量约束仍为2022年四川省体育产业总产出2,170.80亿元。界面中的演示指标不参与这些数量的重新计算。",
        bold_prefix="内容一致性审查",
    )

    remove_paragraph(old_note)
    remove_paragraph(old_caption)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    result = Document(output)
    final_counts = {
        "paragraphs": len(result.paragraphs),
        "tables": len(result.tables),
        "inline_shapes": len(result.inline_shapes),
        "sections": len(result.sections),
    }
    if final_counts["tables"] != original_counts["tables"] or final_counts["sections"] != original_counts["sections"]:
        raise RuntimeError("Table or section count changed unexpectedly")
    if final_counts["inline_shapes"] != original_counts["inline_shapes"] + len(SCREENSHOTS):
        raise RuntimeError("Unexpected number of embedded screenshots")

    all_text = "\n".join(p.text for p in result.paragraphs)
    if "截图预留" in all_text or "最终提交前可将图中九个演示位替换" in all_text:
        raise RuntimeError("Figure 4-2 placeholder language remains in the output")
    table_text = "\n".join(cell.text for table in result.tables for row in table.rows for cell in row.cells)
    if "Backend 269 pass" in table_text or "Frontend 32/32" in table_text:
        raise RuntimeError("Stale quality-gate metrics remain in the output")
    if "Backend：聚焦48/48" not in table_text or "Frontend：34/34+build" not in table_text:
        raise RuntimeError("Audited quality-gate metrics are missing from the output")

    return {
        "source": str(source),
        "output": str(output),
        "source_sha256": sha256(source),
        "output_sha256": sha256(output),
        "original_counts": original_counts,
        "final_counts": final_counts,
        "manual_toc_entries_updated": toc_updates_applied,
        "quality_gate_rows_updated": quality_gate_rows_updated,
        "placeholder_language_removed": True,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--artifact-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    raw_dir = args.artifact_dir / "raw"
    processed_dir = args.artifact_dir / "processed"

    screenshot_manifest = []
    for item in SCREENSHOTS:
        raw_path = raw_dir / f"{item['id']}.png"
        if not raw_path.exists():
            raise FileNotFoundError(raw_path)
        screenshot_manifest.append(
            create_framed_screenshot(raw_path, processed_dir / raw_path.name, item)
        )

    report_manifest = update_report(args.source, args.output, processed_dir)
    manifest = {
        "screenshots": screenshot_manifest,
        "report": report_manifest,
    }
    manifest_path = args.artifact_dir / "report_build_manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
